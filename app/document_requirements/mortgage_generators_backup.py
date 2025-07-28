# document_requirements/mortgage_generators.py - Complete Mortgage Document Generation

from django.utils import timezone
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import logging

logger = logging.getLogger(__name__)


class MortgageChargeGenerator:
    """Generator class for Mortgage and Charge documents with all helper methods"""

    @classmethod
    def generate_document(cls, requirement):
        """Main entry point for generating Mortgage and Charge document"""
        try:
            # Import docx libraries
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
            except ImportError as e:
                logger.error(f"python-docx library not installed: {str(e)}")
                return None

            # Get application data
            application = requirement.application
            context = cls._get_mortgage_charge_context(requirement)
            logger.info(f"Processing Mortgage and Charge for application {application.id}")

            # Create Word document with professional styling
            doc = Document()

            # Set professional margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
                section.page_height = Inches(11.69)  # A4 height
                section.page_width = Inches(8.27)  # A4 width

            # Add all sections
            cls._add_mortgage_header(doc, context)
            cls._add_contents_table(doc, context)
            cls._add_party_details(doc, context)
            cls._add_whereas_clauses(doc, context)
            cls._add_interpretation_section(doc, context)
            cls._add_covenant_to_pay_section(doc, context)
            cls._add_grant_of_security_section(doc, context)
            cls._add_provisions_as_to_security_section(doc, context)
            cls._add_perfection_of_security_section(doc, context)
            cls._add_representations_warranties_section(doc, context)
            cls._add_undertakings_section(doc, context)
            cls._add_rights_of_enforcement_section(doc, context)
            cls._add_enforcement_of_security_section(doc, context)
            cls._add_receivers_section(doc, context)
            cls._add_set_off_section(doc, context)
            cls._add_release_of_security_section(doc, context)
            cls._add_waiver_of_defences_section(doc, context)
            cls._add_new_account_section(doc, context)
            cls._add_application_of_proceeds_section(doc, context)
            cls._add_suspense_account_section(doc, context)
            cls._add_power_of_attorney_section(doc, context)
            cls._add_expenses_and_indemnity_section(doc, context)
            cls._add_currencies_section(doc, context)
            cls._add_transfers_section(doc, context)
            cls._add_notices_section(doc, context)
            cls._add_miscellaneous_section(doc, context)
            cls._add_counterparts_section(doc, context)
            cls._add_governing_law_section(doc, context)
            cls._add_jurisdiction_section(doc, context)
            cls._add_all_schedules(doc, context)
            cls._add_execution_page(doc, context)

            # Save to BytesIO
            result = BytesIO()
            doc.save(result)
            result.seek(0)
            logger.info("Mortgage and Charge document created successfully")
            return result

        except Exception as e:
            logger.error(f"Error generating Mortgage and Charge document: {str(e)}")
            import traceback
            logger.error(f"Full traceback: {traceback.format_exc()}")
            return None

    @classmethod
    def _get_mortgage_charge_context(cls, requirement):
        """Get context data specifically for Mortgage and Charge"""
        application = requirement.application

        # Get today's date formatted
        today = timezone.now().strftime('%d %B %Y')

        context = {
            'execution_date': today,
            'chargor_name': '',
            'lender_name': '',
            'loan_agreement_date': '',
            'folio_number': '',
            'county': '',
            'property_address': '',
            'application': application,
        }

        return context

    @classmethod
    def _add_mortgage_header(cls, doc, context):
        """Add professional document header"""
        # Main title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run('MORTGAGE AND CHARGE')
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.name = 'Times New Roman'
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        title_para.paragraph_format.space_after = Pt(24)
        title_para.paragraph_format.space_before = Pt(12)

    @classmethod
    def _add_contents_table(cls, doc, context):
        """Add contents table matching the original"""
        contents_para = doc.add_paragraph()
        contents_run = contents_para.add_run('CONTENTS')
        contents_run.bold = True
        contents_run.font.size = Pt(14)
        contents_run.font.name = 'Times New Roman'
        contents_para.paragraph_format.space_after = Pt(12)

        # Create contents table
        contents_table = doc.add_table(rows=1, cols=2)
        contents_table.style = 'Table Grid'
        contents_table.columns[0].width = Inches(5.5)
        contents_table.columns[1].width = Inches(1.5)

        # Remove default row
        contents_table._element.remove(contents_table.rows[0]._element)

        # Add header row
        header_row = contents_table.add_row()
        header_row.cells[0].paragraphs[0].add_run('Clause').bold = True
        header_row.cells[1].paragraphs[0].add_run('Page').bold = True

        # Add all contents entries
        contents_entries = [
            ('1 INTERPRETATION', '1'),
            ('2 COVENANT TO PAY', '7'),
            ('3 GRANT OF SECURITY', '8'),
            ('4 PROVISIONS AS TO SECURITY', '9'),
            ('5 PERFECTION OF SECURITY', '10'),
            ('6 REPRESENTATIONS AND WARRANTIES', '11'),
            ('7 UNDERTAKINGS', '13'),
            ('8 RIGHTS OF ENFORCEMENT', '17'),
            ('9 ENFORCEMENT OF SECURITY', '18'),
            ('10 RECEIVERS', '20'),
            ('11 SET-OFF', '23'),
            ('12 RELEASE OF SECURITY', '23'),
            ('13 WAIVER OF DEFENCES', '24'),
            ('14 NEW ACCOUNT', '24'),
            ('15 APPLICATION OF PROCEEDS', '25'),
            ('16 SUSPENSE ACCOUNT', '25'),
            ('17 POWER OF ATTORNEY', '25'),
            ('18 EXPENSES AND INDEMNITY', '26'),
            ('19 CURRENCIES', '26'),
            ('20 TRANSFERS', '26'),
            ('21 NOTICES', '27'),
            ('22 MISCELLANEOUS', '28'),
            ('23 COUNTERPARTS', '28'),
            ('24 GOVERNING LAW', '28'),
            ('25 JURISDICTION', '28'),
            ('Schedule 1', '30'),
            ('Part I Scheduled Property', '30'),
            ('Part II FORM 52', '31'),
            ('Schedule 2', '32'),
            ('Part I Occupational Leases', '32'),
            ('Schedule 3', '33'),
            ('Part I Form of notice relating to assigned Occupational Leases', '33'),
            ('Part II Acknowledgement of Notice of Assignment of Occupational Leases', '35'),
            ('Schedule 4', '36'),
            ('Part I Form of notice relating to Insurances', '36'),
            ('Part II Acknowledgement of Notice of Assignment of Insurances', '37'),
        ]

        for clause_title, page_num in contents_entries:
            row = contents_table.add_row()
            row.cells[0].paragraphs[0].add_run(clause_title)
            row.cells[1].paragraphs[0].add_run(page_num)

        doc.add_paragraph().paragraph_format.space_after = Pt(18)

    @classmethod
    def _add_party_details(cls, doc, context):
        """Add party details section with professional formatting"""
        # THIS DEED paragraph
        deed_para = doc.add_paragraph()
        deed_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        deed_run = deed_para.add_run('THIS DEED ')
        deed_run.bold = True
        deed_run.font.size = Pt(12)
        deed_run.font.name = 'Times New Roman'
        deed_para.add_run(f"is dated {context.get('execution_date', '___')} 202__ and is made")
        deed_para.paragraph_format.space_after = Pt(12)
        deed_para.paragraph_format.line_spacing = 1.15

        # BETWEEN section
        between_para = doc.add_paragraph()
        between_run = between_para.add_run('BETWEEN:')
        between_run.bold = True
        between_run.font.size = Pt(12)
        between_run.font.name = 'Times New Roman'
        between_para.paragraph_format.space_after = Pt(8)

        # Party 1
        party1_para = doc.add_paragraph()
        party1_para.paragraph_format.left_indent = Inches(0.5)
        party1_para.add_run('1. ').bold = True
        party1_para.add_run(f"{context.get('chargor_name', '_' * 40)} ")
        chargor_bold = party1_para.add_run('(the "Chargor")')
        chargor_bold.bold = True
        party1_para.add_run('; and')
        party1_para.paragraph_format.space_after = Pt(8)

        # Party 2
        party2_para = doc.add_paragraph()
        party2_para.paragraph_format.left_indent = Inches(0.5)
        party2_para.add_run('2. ').bold = True
        party2_para.add_run(f"{context.get('lender_name', '_' * 40)} ")
        lender_bold = party2_para.add_run('(the "Lender")')
        lender_bold.bold = True
        party2_para.add_run('.')
        party2_para.paragraph_format.space_after = Pt(12)

    @classmethod
    def _add_whereas_clauses(cls, doc, context):
        """Add WHEREAS clauses with professional formatting"""
        whereas_para = doc.add_paragraph()
        whereas_run = whereas_para.add_run('WHEREAS:')
        whereas_run.bold = True
        whereas_run.font.size = Pt(12)
        whereas_run.font.name = 'Times New Roman'
        whereas_para.paragraph_format.space_after = Pt(8)

        # Whereas A
        whereas_a = doc.add_paragraph()
        whereas_a.paragraph_format.left_indent = Inches(0.5)
        whereas_a.paragraph_format.first_line_indent = Inches(-0.25)
        whereas_a_run = whereas_a.add_run('A. ')
        whereas_a_run.bold = True
        whereas_a.add_run(
            'The Chargor is now or may after the date of this Deed become indebted to the Lender as principal, surety or otherwise.')
        whereas_a.paragraph_format.space_after = Pt(8)
        whereas_a.paragraph_format.line_spacing = 1.15

        # Whereas B
        whereas_b = doc.add_paragraph()
        whereas_b.paragraph_format.left_indent = Inches(0.5)
        whereas_b.paragraph_format.first_line_indent = Inches(-0.25)
        whereas_b_run = whereas_b.add_run('B. ')
        whereas_b_run.bold = True
        whereas_b.add_run(
            'The parties to this Deed have agreed and it is hereby intended that this Deed shall secure the payment, discharge and performance of the Secured Liabilities.')
        whereas_b.paragraph_format.space_after = Pt(12)
        whereas_b.paragraph_format.line_spacing = 1.15

        # IT IS AGREED
        agreed_para = doc.add_paragraph()
        agreed_run = agreed_para.add_run('IT IS AGREED ')
        agreed_run.bold = True
        agreed_run.font.size = Pt(12)
        agreed_para.add_run('as follows:')
        agreed_para.paragraph_format.space_after = Pt(16)

    @classmethod
    def _add_interpretation_section(cls, doc, context):
        """Add Clause 1 - INTERPRETATION with all definitions"""
        # Main clause heading
        cls._add_main_clause_header(doc, '1', 'INTERPRETATION')

        # 1.1 Definitions
        cls._add_subsection_header(doc, '1.1', 'Definitions')

        # Definitions intro
        intro_para = doc.add_paragraph()
        intro_para.paragraph_format.left_indent = Inches(0.5)
        intro_para.add_run('In this Deed:')
        intro_para.paragraph_format.space_after = Pt(8)

        # All definitions from the original document
        definitions = [
            ('"Act"', 'means the Land and Conveyancing Law Reform Act 2009;'),

            ('"Ancillary Covenants"', '''means all covenants, undertakings, guarantees, bonds, warranties, indemnities and other agreements in respect of:

                (a) the design, construction, fit-out or maintenance of any building, structure or erection now or after the date of this Deed on the Scheduled Property (or any part of the Scheduled Property); or
                
                (b) of any roads, footpaths or utilities for services now or after the date of this Deed abutting or serving the Scheduled Property (or any part of the Scheduled Property or
                
                (c) the taking in charge of the Scheduled Property or the paying of any charge or levy in respect of the Scheduled Property, the benefit of which is now or after the date of this Deed vested in the Chargor;'''),

            ('"Business Day"',
             'means a day (other than a Saturday or Sunday) on which banks in Dublin are open for general banking business;'),

            ('"Charged Assets"',
             'means all present and future assets, rights and property of the Chargor the subject of any security created or expressed or intended to be created by or pursuant to this Deed and in the Prescribed Form Charge and any reference to the "Charged Assets" includes a reference to any of them;'),

            ('"Compensation Rights"',
             'means all present and future rights of the Chargor to be paid or receive compensation under any statute or by reason of any compulsory acquisition, requisition or other exercise of compulsory powers in relation to the Scheduled Property (or any part of the Scheduled Property) or arising on any refusal, withdrawal or modification of planning permission or approval relative to the Scheduled Property or any control or limitation imposed upon or affecting the use of the Scheduled Property (or any part of the Scheduled Property);'),

            ('"Default Rate"',
             'means the rate per annum payable in respect of any overdue amount as determined by the Lender from time to time;'),

            ('"Development"', 'shall have the meaning ascribed to development under the Planning Acts;'),

            ('"Environment"',
             'means any land (including surface and sub-surface soil, the sea bed and any natural or man-made structures), water (including rivers, lakes (man-made or natural), canals, the ocean (whether within or without territorial waters), ground waters and waters in drains and sewers) and air (including air within buildings and other natural and man-made structures) above or below ground;'),

            ('"Environmental Laws"',
             'means all and any applicable laws, including common law, statute, bye-law and subordinate legislation, regulations, codes of practice, circulars, and directives and judgments and decisions (whether in Ireland or elsewhere and whether or not having the force of law), including notices, orders or circulars, of any court or authority competent to make such judgment or decision compliance with which is mandatory for the Chargor in any jurisdiction with regard to Environmental Matters and protection of the Environment including the Public Health (Ireland) Act 1878, the Air Pollution Act 1987 and 2011, the Local Government (Water Pollution) Acts 1977 to 2007, the Fisheries Acts 1959 to 2007, the Dangerous Substances Acts 1972 and 1979, the Litter Pollution Acts 1997 to 2009, the Safety, Health and Welfare at Work Act 2005 to 2014, the Safety in Industry Act 1955 to 1987, the Planning Acts, the Environmental Protection Agency Act 1992 to 2011, the Waste Management Acts 1996 to 2011, the European Communities Act 1972 to 2012 and all regulations, bye-laws, orders, decisions and codes made thereunder;'),

            ('"Environmental Matters"', '''means any matter arising out of, relating to, or resulting from:

                (a) the pollution or protection of the Environment;
                
                (b) harm to the health of humans, animals or plants including laws relating to public and workers' health and safety;
                
                (c) emissions, discharges or releases into the Environment of chemicals or any other pollutants or contaminants or industrial, radioactive, dangerous, toxic or hazardous substances or waste (whether in solid, semi-solid, liquid or gaseous form and including noises and genetically modified organisms); or
                
                (d) the manufacture, processing, use, treatment, storage, distribution, disposal, transport or handling of the substances or wastes described in sub-paragraph (c);'''),

            ('"Environmental Permit"',
             'means all and any permits, licences, consents, approvals, certificates, qualifications, specifications, registrations and other authorisations (including any conditions which attach to any of the foregoing) and the filing of all notifications, reports and assessments required by any Environmental Law;'),

            ('"Event of Default"', '''means:

                (a) in the case of any of the Secured Liabilities which are repayable on demand, the making by the Lender of a demand for repayment in relation to the Secured Liabilities; or
                
                (b) any event, howsoever described, as specified in any Finance Document the occurrence of which entitles the Lender to accelerate or demand early repayment of the Secured Liabilities or, in the case of any part of the Secured Liabilities, which is at such time contingent, to call for the delivery of cash collateral in respect of the Secured Liabilities or any failure by the Chargor to pay or repay on demand all or any of the Secured Liabilities which are so payable; or
                
                (c) any Event of Default referred to in clause 8 (Rights of Enforcement);'''),

            ('"Finance Document"', '''means:

                (a) the Loan Agreement; and
                
                (b) each loan agreement, guarantee, indemnity, counter-indemnity, bond, guarantee, hiring agreement, leasing agreement, hire purchase agreement, credit sale agreement, factoring agreement, invoice discounting agreement, debt purchase agreement, bill of exchange, promissory note, or any other agreement of any nature or kind (both present and future) containing or evidencing the terms upon which or under which any financing, credit or other facilities have been made available by the Lender to the Chargor or to some other party at the request and on behalf of the Chargor as the case may be or any instrument or document (including this Deed) which creates or evidences any Security Interest and which is entered into in connection therewith;'''),

            ('"Financial Indebtedness"', '''means any indebtedness or liability in respect of:

                (a) monies borrowed or raised and debit balances at banks or other financial institutions;
                
                (b) any debenture, bond, any subordinated loan note, loan stock or other security;
                
                (c) any acceptance credit;
                
                (d) receivables sold or discounted (otherwise than on a non-recourse basis);
                
                (e) the acquisition cost of any asset to the extent payable before or after the time of acquisition or possession by the party liable where the advance or deferred payment is arranged primarily as a method of raising finance or financing the acquisition of that asset;
                
                (f) leases entered into primarily as a method of raising finance or financing the acquisition of assets leased;
                
                (g) currency, interest rate swaps or other derivatives;
                
                (h) amounts raised under any other transaction having the commercial effect of the borrowing or raising money save for any trade credit terms arising in the ordinary course of business; or
                
                (i) any guarantee, letter of credit, indemnity or similar assurance against financial loss of any person save in respect of any indebtedness already included in sub-clauses (a) to (h) inclusive;'''),

            ('"Insurance Proceeds"',
             'means all proceeds paid or payable to the Chargor under or in connection with the Insurances;'),

            ('"Insurances"',
             'means in relation to the Charged Assets, all contracts and policies of insurance and re-insurance of any kind taken out or, as the context requires, to be taken out and maintained by or on behalf of the Chargor or in which the Chargor has an interest including the contracts and policies existing at the date of this Deed as more particularly listed in Part III of Schedule 2 (Insurances) and the debts and benefits represented by such contracts and/or policies;'),

            ('"Loan Agreement"',
             f'means the loan agreement dated {context.get("loan_agreement_date", "_" * 20)} entered into between the Chargor and the Lender (as the same may be amended, restated, supplemented, varied, extended, novated and/or replaced from time to time);'),

            ('"Occupational Leases"',
             'means the leases, licenses, agreements for lease, all licences and other agreements for the occupation, possession or use of all or any part or parts of the Scheduled Property including those listed in Part I of Schedule 2 (Occupational Leases) subject to which the interest of the Chargor in the Scheduled Property is now or from time to time after the date of this Deed held and "Occupational Lease" includes a reference to any of them;'),

            ('"Planning Acts"',
             'means all laws (whether criminal, civil or administrative) including common law, statute, statutory instruments, directives, regulations, bye-laws, orders, codes, judgments and other legal measures having the force of law concerning planning matters including the Planning and Development Acts 2000 to 2021, the Building Control Acts 1990 to 2014, the Local Government (Planning and Development) Acts 1963 to 1999 and the Fire Services Acts 1981 to 2003 and any regulations issued pursuant to such Acts and any extant order or regulation made or confirmed under any of them;'),

            ('"Prescribed Form"',
             'means the form of charge set out in Part II of Schedule 1 (Form 52) or such other form of charge as, in the opinion of the Lender, may be required at law to charge registered land;'),

            ('"Prescribed Form Charge"',
             'means a charge in the Prescribed Form entered into or to be entered into by the Chargor in favour of the Lender;'),

            ('"Receiver"',
             'means any one or more receivers and/or managers appointed by the Lender in respect of the Chargor or over all or any of its property, assets or undertaking (whether pursuant to this Deed, at law or otherwise) which shall, where the context so admits, include the plural and any replacement or substitute receiver and/or manager;'),

            ('"Rental Income"',
             'means all amounts payable or paid to or for the benefit of the Chargor pursuant to, or in connection with or arising under any Occupational Lease and including, for the avoidance of doubt, all rents, licence fees, premiums, key monies, mesne profits and any interest payable in respect of any of the foregoing;'),

            ('"Scheduled Property"',
             'means the property described in Part I of Schedule 1 (Scheduled Property) and the property described in the Prescribed Form Charge and (by way of extension of section 71 of the Act) all rights, liberties, powers, easements, quasi easements and appurtenances (in each case of whatever nature) attached or appurtenant to the Scheduled Property and all buildings, erections, fixtures, fittings (including trade fixtures and fittings) from time to time in or on the Scheduled Property and the full benefit of all warranties and maintenance contracts for any of the same;'),

            ('"Secured Liabilities"',
             'means all present and future obligations and liabilities of the Chargor to the Lender to include, without limitation, all liabilities due and owing from time to time to the Lender in respect of principal, interest (including any interest which has been rolled up or capitalised and default interest), break-costs, discount, commission, fees and expenses (including all costs and expenses of and incidental to the appointment of a Receiver and the exercise of all or any of his powers) and whether such liabilities are actual or contingent, whether owed solely or jointly with any other person, whether as principal or as surety or in any other capacity whatsoever and in any currency and on any current or other account, whether under any Finance Document or otherwise in any manner whatsoever and any reference to the "Secured Liabilities" includes a reference to any of them;'),

            ('"Security"',
             'means the security constituted or intended to be constituted by this Deed and each Prescribed Form Charge;'),

            ('"Security Interest"',
             'means any mortgage, charge (fixed or floating), pledge, lien, assignment, hypothecation, title retention, preferential right, trust arrangement or other security arrangement or agreement (including the deposit of monies or property with a person with the intention of affording such person a right of set-off or lien) and any other agreement or arrangement having a similar effect;'),

            ('"Security Period"',
             'means the period commencing on the date of execution of this Deed and terminating upon the date on which the Lender notifies the Chargor that the Secured Liabilities have been irrevocably and unconditionally paid and discharged in full and all the commitments of the Lender cancelled; and'),

            ('"Work-in-Progress"',
             'means all construction and other works carried out or in the process of being carried out from time to time on or in connection with the Scheduled Property including all site preparatory works, all demolition and site clearance works and all materials or goods which are intended for inclusion in, or are in the process of being included in, any of the foregoing works.'),
        ]

        for term, definition in definitions:
            cls._add_definition(doc, term, definition)

        # Add remaining subsections
        cls._add_subsection_header(doc, '1.2', 'Successors and Assigns')
        cls._add_subsection_content(doc,
                                    'References to the "Lender", the "Chargor" or any other person include references to their successors and permitted assignees, transferees, novates, personal representatives or substitutes (as appropriate) whether immediate or derivative or the acquisition of all or part of the undertaking of the Lender by any other person.')

        cls._add_subsection_header(doc, '1.3', 'Headings')
        cls._add_subsection_content(doc,
                                    'Clause headings and the contents page are inserted for convenience of reference only and shall be ignored in the interpretation of this Deed.')

        cls._add_subsection_header(doc, '1.4', 'Construction')
        cls._add_construction_subsections(doc)

    @classmethod
    def _add_covenant_to_pay_section(cls, doc, context):
        """Add Clause 2 - COVENANT TO PAY"""
        cls._add_main_clause_header(doc, '2', 'COVENANT TO PAY')

        # 2.1 Covenant to pay and discharge
        cls._add_subsection_header(doc, '2.1', 'Covenant to pay and discharge')

        # 2.1.1
        cls._add_numbered_subsection(doc, '2.1.1',
                                     'The Chargor hereby unconditionally and irrevocably covenants with the Lender that he shall, on demand by the Lender (if expressed to be so payable) or on such terms as may otherwise be agreed in writing between the Chargor and the Lender, pay, discharge and perform the Secured Liabilities. In the absence of any formal agreement to the contrary, the Chargor acknowledges and confirms that any liabilities in respect of the Secured Liabilities shall be due and payable to the Lender on demand.')

        # 2.1.2
        cls._add_numbered_subsection(doc, '2.1.2',
                                     'The making of one demand shall not preclude the Lender from making any further demands.')

        # 2.2 Interest
        cls._add_subsection_header(doc, '2.2', 'Interest')
        cls._add_subsection_content(doc,
                                    'The Chargor shall pay interest on each amount demanded of her under this Deed from the due date for payment until actual payment (after as well as before judgment) at the Default Rate.')

        # 2.3 Payment free of deduction
        cls._add_subsection_header(doc, '2.3', 'Payment free of deduction')
        cls._add_subsection_content(doc,
                                    'All payments to be made under this Deed by the Chargor shall be made free and clear of and without deduction for or on account of any set‑off, counterclaim or any present or future taxes, levies, imposts, duties, charges, fees, deductions or withholdings of any nature whatsoever. If the Chargor shall at any time be compelled by law to make any deduction or withholding from any payment to be made under this Deed, the Chargor will concurrently pay to the Lender such additional amount as will result in payment to the Lender of the full amount which would have been received had such deduction or withholding not been made and will, on request, supply to the Lender all appropriate documentation (in each case in a form and substance satisfactory to the Lender) evidencing that the Chargor has duly accounted to the relevant authority for any such deduction or withholding.')

        # 2.4 Evidence and calculation
        cls._add_subsection_header(doc, '2.4', 'Evidence and calculation')
        cls._add_subsection_content(doc,
                                    'Any certificate or determination by the Lender as to the amount of the Secured Liabilities shall, in the absence of manifest error or fraud, be conclusive and binding on the Chargor for all purposes.')

        # 2.5 Obligations Unconditional
        cls._add_subsection_header(doc, '2.5', 'Obligations Unconditional')
        cls._add_subsection_content(doc,
                                    'The obligations of the Chargor under clause 2.1 (Covenant to pay and discharge) are unconditional and neither the provisions of this Deed nor the obligations of the Chargor will be affected by the occurrence or existence at any time of any of the following events or circumstances or by any person\'s knowledge or lack of knowledge as to any such matter:')

        # Add numbered list
        unconditional_events = [
            'any person\'s insolvency or lack of capacity, power or authority;',
            'any unenforceability, illegality or invalidity of any obligation of any person;',
            'any change in the constitution, membership, ownership, legal form, name or status of any person;',
            'the making, amendment or termination of any other deed or agreement;',
            'any amendment, novation, re-statement or substitution of, or any supplement to, any other deed or agreement;',
            'any increase or reduction in the amount of any person\'s indebtedness or any alteration of any term, condition or arrangement in respect of any person\'s indebtedness;',
            '''any person taking or omitting to take any steps in relation to:

    (a) the Chargor or any other person;

    (b) any of the Secured Liabilities;

    (c) any Security, guarantee or other financial support in respect of any indebtedness; and/or

    (d) any other asset; or''',
            'anything else which, although it could affect the liability of a surety, would not affect the liability of a principal debtor.'
        ]

        for i, event in enumerate(unconditional_events, 1):
            cls._add_numbered_list_item(doc, str(i), event)

    @classmethod
    def _add_grant_of_security_section(cls, doc, context):
        """Add Clause 3 - GRANT OF SECURITY"""
        cls._add_main_clause_header(doc, '3', 'GRANT OF SECURITY')

        # 3.1 Fixed Charges
        cls._add_subsection_header(doc, '3.1', 'Fixed Charges')
        cls._add_subsection_content(doc,
                                    'Subject to clause 12.1 (Release of Security), the Chargor for good valuable consideration and as continuing security for the payment, discharge and performance of the Secured Liabilities and as legal and beneficial owner (and where applicable as registered owner or, as the case may be, the person entitled to be registered as owner) hereby:')

        # Add numbered charges
        charges = [
            'charges in favour of the Lender, by way of first fixed charge, the Scheduled Property and hereby assents to the registration of such charge as a burden on the said property;',
            'charges in favour of the Lender by way of first fixed charge, all other present and future estate, right, title or interest of the Chargor in any freehold or leasehold property or any lands hereditaments or premises comprised in the Scheduled Property together in all cases with all buildings, fixtures and fittings (including trade fixtures) from time to time on such freehold or leasehold property or such lands hereditaments or premises comprised in the Scheduled Property;',
            'charges in favour of the Lender by way of first fixed charge, all other present and future estate, right, title or interest (but not its obligations) in and to any Development being undertaken by the Chargor on all or any part of the Scheduled Property and the Work-in-Progress;',
            'charges in favour of the Lender by way of first fixed charge, all its right, title, benefits, entitlements, remedies and interests in and payments under (but not the burden of) any Ancillary Covenants to hold same absolutely;',
            'charges in favour of the Lender by way of first fixed charge, all present and future rights, title and interest of the Chargor in and to the Occupational Leases, the Rental Income (and the benefit of any guarantee and/or security given in connection with the Occupational Leases and/or the Rental Income) and any other present and future lease or licence of the Scheduled Property and any rights or benefits deriving from the Occupational Leases, the Rental Income (and the benefit of any guarantee and/or security given in connection with the Occupational Leases and/or the Rental Income) and any other present and future lease or licence of the Scheduled Property;',
            'charges in favour of the Lender by way of first fixed charge, (to the extent not effectively assigned or otherwise charged under this Deed) the Insurances and the Insurance Proceeds; and',
            'charges in favour of the Lender by way of first fixed charge all Compensation Rights,'
        ]

        for i, charge in enumerate(charges, 1):
            cls._add_numbered_list_item(doc, f'3.1.{i}', charge)

        cls._add_subsection_content(doc,
                                    'and in the case of all of the charges and the assignments at clauses 3.1.1 to 3.1.6, these shall be deemed to include all damages, compensation, remuneration, profit, rent or other monies which the Chargor may now or at any time derive therefrom.')

        # 3.2 Assignments
        cls._add_subsection_header(doc, '3.2', 'Assignments')
        cls._add_subsection_content(doc,
                                    'Subject to clause 12.1 (Release of Security), the Chargor for good valuable consideration and as continuing security for the payment, discharge and performance of the Secured Liabilities and as legal and beneficial owner hereby assigns and agrees to assign absolutely to the Lender by way of first fixed security:')

        assignments = [
            'all his right, title, benefits, entitlements, remedies and interests in and payments under (but not the burden of) any Ancillary Covenants to hold same absolutely;',
            'all of his right, title and interest in the Occupational Leases, the Rental Income (and the benefit of any guarantee and/or security given in connection with the Occupational Leases and/or the Rental Income) and any other present and future lease or licence of the Scheduled Property;',
            'all of his right, title and interest in the Insurances and the Insurance Proceeds;',
            'all of his right, title and interest in all the Chargor\'s present and future goodwill in so far as it relates to the Scheduled Property; and',
            'all Compensation Rights'
        ]

        for i, assignment in enumerate(assignments, 1):
            cls._add_numbered_list_item(doc, f'3.2.{i}', assignment)

        cls._add_subsection_content(doc,
                                    'and in the case of all of the assignments at clauses 3.2.1 to 3.2.4 these shall be deemed to include all damages, compensation, remuneration, profit, rent or other monies which the Chargor may now or at any time derive therefrom.')

    @classmethod
    def _add_provisions_as_to_security_section(cls, doc, context):
        """Add Clause 4 - PROVISIONS AS TO SECURITY"""
        cls._add_main_clause_header(doc, '4', 'PROVISIONS AS TO SECURITY')

        # 4.1 Continuing security
        cls._add_subsection_header(doc, '4.1', 'Continuing security')

        cls._add_numbered_subsection(doc, '4.1.1',
                                     'The Security is a continuing security over all present and future assets and undertaking of the Chargor intended to be charged and assigned under this Deed and will remain in full force and effect as a continuing security for the payment and discharge of the Secured Liabilities until released or discharged by the Lender.')

        cls._add_numbered_subsection(doc, '4.1.2',
                                     'No part of the Security will be considered satisfied or discharged by any intermediate payment, discharge or satisfaction of the whole or any part of the Secured Liabilities.')

        cls._add_numbered_subsection(doc, '4.1.3',
                                     'If upon the final repayment and satisfaction of the Secured Liabilities there shall exist any right on the part of the Chargor or any other person to draw funds or otherwise which, if exercised, would or might cause the Chargor to become actually or contingently liable to the Lender whether as principal debtor or as surety for another person, then the Lender will be entitled to retain this Security and all rights, remedies and powers conferred by this Deed and the Charged Assets for so long as the Lender, acting reasonably, deems necessary and in the event that any demand is made by the Lender under any Finance Document the said monies will become due and shall be paid and discharged to the Lender and all provisions of this Deed shall apply accordingly.')

        # 4.2 Additional security
        cls._add_subsection_header(doc, '4.2', 'Additional security')
        cls._add_subsection_content(doc,
                                    'This Deed is in addition to and is not prejudiced by any other security now or after the date of this Deed held by the Lender for the Secured Liabilities.')

        # 4.3 Non-competition
        cls._add_subsection_header(doc, '4.3', 'Non-competition')
        cls._add_subsection_content(doc,
                                    'Until the Security has been discharged, the Chargor will not, after a claim has been made or by virtue of any payment or performance by the Chargor of the Secured Liabilities:')

        non_competition_items = [
            'be subrogated to any rights, security or monies held, received or receivable by the Lender or any person nor be entitled to any right of contribution or indemnity in respect of any payment made or monies received on account of the Secured Liabilities;',
            'claim, rank, prove or vote as creditor of any person; or',
            'receive, claim or have the benefit of any payment, distribution or security from or on account of any person, or exercise any right of set-off as against any person, and'
        ]

        for i, item in enumerate(non_competition_items, 1):
            cls._add_numbered_list_item(doc, str(i), item)

        cls._add_subsection_content(doc,
                                    'the Chargor shall forthwith pay or transfer to the Lender an amount equal to the amount of any dividend, distribution, contribution or benefit (including any amount set-off) actually received by it and in the meantime shall hold the same in trust for the Lender to the extent required to pay or discharge the Secured Liabilities.')

        # 4.4 Negative pledge
        cls._add_subsection_header(doc, '4.4', 'Negative pledge')
        cls._add_subsection_content(doc,
                                    'The Chargor hereby unconditionally and irrevocably covenants with the Lender that he will not during the Security Period without the prior written consent of the Lender:')

        negative_pledge_items = [
            'create or attempt to create or permit to subsist in favour of any person other than the Lender, any Security Interest in or affecting the Charged Assets or any part of the Charged Assets;',
            'dispose of the Charged Assets or any part of the Charged Assets or attempt or agree so to do; or',
            'permit any variation, waiver or termination of any of the rights attaching to the whole or any part of the Charged Assets.'
        ]

        for i, item in enumerate(negative_pledge_items, 1):
            cls._add_numbered_list_item(doc, str(i), item)

        # 4.5 Consent to Registration
        cls._add_subsection_header(doc, '4.5', 'Consent to Registration')
        cls._add_subsection_content(doc,
                                    'The Chargor hereby irrevocably consents to the registration of all or any of the Security as a burden on the property and assets thereby affected.')

    @classmethod
    def _add_perfection_of_security_section(cls, doc, context):
        """Add Clause 5 - PERFECTION OF SECURITY"""
        cls._add_main_clause_header(doc, '5', 'PERFECTION OF SECURITY')

        # 5.1 Notices of assignment
        cls._add_subsection_header(doc, '5.1', 'Notices of assignment')
        cls._add_subsection_content(doc, 'The Chargor hereby covenants with the Lender that he shall:')

        notices_items = [
            'immediately upon execution of this Deed (or, if later, as soon as possible after the Chargor enters into an Occupational Lease), in respect of each Occupational Lease, deliver a duly completed notice of assignment to each other party to that Occupational Lease in the form set out in Part I of Schedule 3 (Form of notice relating to assigned Occupational Leases) and shall procure that each such party who receives such a notice executes and delivers to the Lender an acknowledgement, in the form set out in Part II of Schedule 3 (Acknowledgement of Notice of Assignment of Occupational Leases) (or in each case, in such other form as the Lender shall agree); and',
            'immediately upon execution of this Deed (or, if later, as soon as possible after the Chargor enters into a policy of insurance), deliver, to each insurance company which has issued a policy of insurance, a duly completed notice of assignment in the form set out in Part I of Schedule 4 (Form of notice Relating to Insurances) and shall procure that each such insurance company executes and delivers to the Lender an acknowledgement, in the form set out in Part II of Schedule 4 (Acknowledgement of Notice of Assignment of Insurances) (or in each case, in such other form as the Lender shall agree).'
        ]

        for i, item in enumerate(notices_items, 1):
            cls._add_numbered_list_item(doc, str(i), item)

        # 5.2 Further assurances
        cls._add_subsection_header(doc, '5.2', 'Further assurances')
        cls._add_subsection_content(doc,
                                    'The Chargor shall, at his own expense, if and when at any time required by the Lender or any Receiver appointed by him, execute such further security deeds or instruments (comprising fixed charges and assignments) including, if required by the Lender (acting reasonably), a deed of confirmation in respect of, inter alia, the continued validity of the Security in each case in favour of or for the benefit of the Lender or any such Receiver and do all such acts and things as may be required to perfect or protect the Security over the Charged Assets (both present and future) or any part of the Charged Assets or to facilitate the realisation of same or the exercise of any right, power or discretion exercisable by the Lender or any such Receiver or any of its delegates or sub-delegates in respect of any Charged Assets.')

        # 5.3 Scheduled Property
        cls._add_subsection_header(doc, '5.3', 'Scheduled Property')
        cls._add_subsection_content(doc,
                                    'Without prejudice to the generality of clause 5.2 (Further assurances), the Chargor hereby covenants and undertakes with the Lender that he will, at his own cost, if and when requested by the Lender, execute a Prescribed Form Charge (with such modifications as the Lender may require) in the form attached in Part II of Schedule 1 over all the land which is, or is intended to be, charged by this Deed and which is registered or in the course of being registered in the Land Registry and will provide all appropriate assistance to the Lender to have the same duly registered in the Land Registry as a burden on the land affected by this Deed.')

        # 5.4 Power to remedy
        cls._add_subsection_header(doc, '5.4', 'Power to remedy')
        cls._add_subsection_content(doc,
                                    'If the Chargor fails to comply with any of his obligations under this Deed, the Lender (or its nominee) may, but shall not be obliged to, (at the Chargor\'s expense) take such action as is reasonably necessary to protect the Charged Assets against the consequences of the Chargor\'s non-compliance and to ensure compliance with such obligations.')

    @classmethod
    def _add_representations_warranties_section(cls, doc, context):
        """Add Clause 6 - REPRESENTATIONS AND WARRANTIES"""
        cls._add_main_clause_header(doc, '6', 'REPRESENTATIONS AND WARRANTIES')

        # 6.1 Representations and warranties
        cls._add_subsection_header(doc, '6.1', 'Representations and warranties')
        cls._add_subsection_content(doc, 'The Chargor represents and warrants to the Lender that:')

        representations = [
            'the Chargor is aged 18 or over, is of sound mind and has not been made a ward of court and has not had an enduring power of attorney registered in respect of him;',
            'the Chargor has not entered into initiated or been the subject of an insolvency arrangements within the meaning of the Personal Insolvency Act, 2012;',
            'the Chargor has not been adjudicated a bankrupt nor have any steps been threatened or taken or proceedings started to have him adjudicated a bankrupt or to have a receiver or similar officer appointed over any of his assets in any jurisdiction;',
            'the Chargor has the power to grant this Deed and to perform his obligations under this Deed;',
            'the Chargor has the power and authority to own his assets and to conduct the business which he conducts and proposes to conduct;',
            'this Deed constitutes the Chargor\'s legal, valid binding and enforceable obligations;',
            'neither the granting of this Deed by the Chargor nor the performance by the Chargor of its obligations hereunder will contravene any law, regulation or any agreement to which he is a party or by which he is bound nor will it cause any limitation on any of its powers howsoever imposed to be exceeded;',
            'all authorisations, consents, approvals, resolutions, licences, exemptions, filings or registrations required for the entry into, performance, validity or enforceability of this Deed by the Chargor have been obtained and are in full force and effect;',
            'this Deed creates those Security Interests it purports to create and is not liable to be avoided or otherwise set aside on its liquidation or otherwise;',
            'save for any Security Interest granted in favour of the Lender, he has not granted or created, nor is there outstanding, any Security Interest over the Charged Assets (or any part of the Charged Assets);',
            'all material covenants, restrictions and stipulations affecting the Charged Assets have been observed and performed and no officer or servant of the Chargor has suffered or committed or caused any breach of any such material covenant, restriction or stipulation;',
            'the Chargor is the legal and beneficial owner of the Charged Assets and has delivered to the Lender all documents evidencing its ownership of the Charged Assets including, in the case of the Scheduled Property, all title deeds and documents in connection with the Scheduled Property;',
            'no Development has taken place on or in relation to the Scheduled Property without any requisite permission having been obtained prior to the commencement of such Development and no situation exists in relation to any of the Scheduled Property in respect of which a warning notice or an enforcement notice has been or may be made;',
            'the Chargor has not received or acknowledged notice of any adverse claim by any person in respect of the Charged Assets or any interest in them;',
            'all factual information provided by the Chargor or on the Chargor\'s behalf was true and accurate in all material respects as at the date it was provided or as at the date (if any) at which he is stated to be given;',
            'he has obtained all consents including the consent of any landlord or superior landlord (where necessary) under any lease or agreement for lease under which the Scheduled Property is held, necessary to ensure that no other party to any agreement or arrangement entered into by it (including such landlord or superior landlord) becomes entitled to terminate such agreement or arrangement as a consequence of it entering into this Deed;',
            'he has complied with all Environmental Laws to which it may be subject including the filing of all notifications, reports or assessments required to be filed under such Environmental Laws; and',
            'he has obtained all Environmental Permits required in connection with his assets and has complied with the terms of those Environmental Permits.'
        ]

        for i, rep in enumerate(representations, 1):
            cls._add_numbered_list_item(doc, str(i), rep)

    @classmethod
    def _add_undertakings_section(cls, doc, context):
        """Add Clause 7 - UNDERTAKINGS"""
        cls._add_main_clause_header(doc, '7', 'UNDERTAKINGS')

        # 7.1 General undertakings
        cls._add_subsection_header(doc, '7.1', 'General undertakings')
        cls._add_subsection_content(doc,
                                    'The Chargor hereby irrevocably covenants and undertakes with the Lender that he shall:')

        general_undertakings = [
            'provide the Lender with such information relating to the Charged Assets as the Lender may reasonably require from time to time;',
            'observe and perform all material covenants, requirements and obligations from time to time imposed on, applicable to or otherwise affecting the Charged Assets and/or the use, ownership, occupation, possession, operation, repair, maintenance or other enjoyment or exploitation of the Charged Assets whether imposed by statute, law or regulation, contract, lease, licence, grant or otherwise;',
            'comply with all laws, statutes and regulations (including those relating to the payment of taxes and the Environment) which are applicable to it and obtain, effect, comply with and maintain in full force and effect all registrations, licences, consents, authorisations and exemptions required for the conduct of its business and the performance, validity and enforceability of this Deed and any document entered into pursuant to this Deed and generally do all other acts and things (including the taking of legal proceedings) necessary or desirable to maintain, defend or preserve its right, title and interest to and in the Charged Assets without infringement by any third party;',
            'not without the prior written consent of the Lender, enter into any onerous or restrictive obligations affecting any of its real property or agree to any rent review (the result of which would mean that it would pay an amount higher than market rent) relating to any interest in any of the Charged Assets;',
            'not without the prior written consent of the Lender, enter into any agreement affecting any of its real property or agree to any rent review (the result of which would mean that it would receive an amount less than market rent) relating to any interest in any of the Charged Assets;',
            'notify the Lender immediately if any distress or execution is levied or enforced against it or any of its assets or any third party debt order or freezing order is made and served on it;',
            'notify the Lender immediately if any steps (including the making of an application without the giving of any notice) are taken by any person (including the Chargor) in relation to receivership, bankruptcy or insolvency arrangements within the meaning of the Personal Insolvency Act, 2012 or any analogous step or procedure is taken in any jurisdiction;',
            'not without the prior written consent of the Lender, do or omit to do anything which might result in any Charged Asset which currently is or becomes a registered right to lapse or which might allow or permit a third party to obtain a revocation of any such registered right; and',
            'not do or allow to be done any act which could, in the reasonable opinion of the Lender, have a material adverse effect on the value of any Charged Asset.'
        ]

        for i, undertaking in enumerate(general_undertakings, 1):
            cls._add_numbered_list_item(doc, str(i), undertaking)

        # 7.2 Insurance
        cls._add_subsection_header(doc, '7.2', 'Insurance')
        cls._add_subsection_content(doc,
                                    'The Chargor hereby irrevocably covenants and undertakes with the Lender that it shall:')

        insurance_undertakings = [
            'insure and keep insured all of the Charged Assets, including its buildings, structures in such amounts, against such risks and with such insurance companies as are acceptable to the Lender including risks such as damage by fire, flood, explosion, riot, civil commotion, loss and liability imposed by law as owner or occupier of any property for damages and such other insurable risks as are commonly insured against from time to time in an amount equal to the full reinstatement or replacement cost (with adequate provision to cover professional fees and loss of rent for a minimum period of three (3) years) and shall procure that the interest of the Lender is endorsed and/or noted on the policy as sole loss payee or with such other interest as the Lender may otherwise direct or agree;',
            'upon execution of this Deed, deliver to the Lender all policies or cover notes (including renewal cover notes or new policies) or other evidence acceptable to the Lender of the insurance required to be maintained in accordance with the provisions of clause 7.2.1;',
            'duly and promptly pay or cause to be paid all premiums and other sums of money payable for maintaining any such insurance as aforesaid;',
            'apply any Insurance Proceeds in making good the loss or damage in respect of which such monies were received or, at the sole discretion of the Lender, in repayment of the Secured Liabilities;',
            'not do or permit anything to be done in or upon or relating to the Charged Assets or any part of the Charged Assets which may make void or voidable any Insurances;',
            'pending application in accordance with clause 7.2.4 or if requested by the Lender to do so pending a determination as to their application or use by the Lender, hold all Insurance Proceeds on trust for the Lender; and',
            'ensure and procure that all policies covering the Charged Assets shall contain clauses whereby the insurer agrees that the policies will not be cancelled or terminated and will not expire without a minimum of thirty (30) days\' notice in writing to the Lender or provisions to this effect and to the extent from time to time available from the insurers.'
        ]

        for i, undertaking in enumerate(insurance_undertakings, 1):
            cls._add_numbered_list_item(doc, f'7.2.{i}', undertaking)

        # 7.3 Scheduled Property
        cls._add_subsection_header(doc, '7.3', 'Scheduled Property')
        cls._add_subsection_content(doc,
                                    'The Chargor hereby irrevocably covenants and undertakes with the Lender that it shall:')

        property_undertakings = [
            'not, without the prior written consent of the Lender, part with possession or occupation of the Scheduled Property (or any part of the Scheduled Property);',
            'comply with, observe and perform all covenants, obligations and conditions relating to the Scheduled Property (including every Occupational Lease, fee farm grant, agreement or other instrument relating to the Scheduled Property) and indemnify the Lender in respect of any breach of those covenants, obligations and conditions;',
            'procure that all Rental Income shall, if so required by the Lender, be paid into such account(s) as the Lender may from time to time specify;',
            'not, without the previous consent in writing of the Lender, which consent shall not be unreasonably withheld, remove or destroy any of the buildings, fixtures, fittings, vehicles, computers and office and other equipment or any structure whatsoever from the Scheduled Property owned by the Chargor unless that property is worn out or rendered unfit for use or unless such removal or destruction shall be with a view to promptly replacing such property by other property of at least equal value or utility;',
            'at all reasonable times permit the Lender and its representatives access to the Scheduled Property from time to time and at the cost of the Chargor to take any action the Lender may consider reasonably necessary or desirable to prevent or remedy any breach of any covenant, stipulation or term of this Deed;',
            'on the expiration of any Occupational Lease or other agreement relating to the Scheduled Property, to enter into negotiations to agree to renew all such Occupational Leases or other agreements on the most favourable terms available and upon receipt of the written consent of the Lender to proceed to renew such Occupational Leases or other agreements;',
            'not, without the prior written consent of the Lender, exercise the statutory powers of leasing or agreeing to lease, granting or agreeing to grant or of accepting or agreeing to accept surrenders conferred by section 18 of the Act in respect of the Scheduled Property or any part of the Scheduled Property or sell, convey, assign, transfer or confer upon any person any contractual licence, right or interest in the Scheduled Property or any part of the Scheduled Property;',
            'not cause or permit any person to become a protected or statutory tenant of the Scheduled Property or any part of the Scheduled Property under the Rent Restrictions Acts 1960 to 1981 or the Housing (Private Rented Dwellings) Acts 1982 and 1983;',
            'in the event of a notice or order given, issued or made to the Chargor affecting the Charged Assets or any part of the Charged Assets or in the event of any proceedings being commenced affecting the Charged Assets, immediately give full particulars of such notice, order or proceedings to the Lender and (without delay and at the cost of the Chargor) take all reasonable or necessary steps to comply with any such notice or order and or make or join with the Lender in making such objections or representations against or in respect of any such notice or order as the Lender shall reasonably require at the sole cost of the Chargor;',
            'notify the Lender promptly of the acquisition of any freehold or leasehold property intended to form part of the Scheduled Property and at any time, if called upon to do so by the Lender and at the Chargor\'s own expense, execute over all or any part of such property a charge, by way of legal charge in favour of the Lender in such form as the Lender may require and in the case of any leasehold property, use its best endeavours to obtain any requisite consent therefor from the lessor;',
            'deposit with the Lender all deeds and documents of title in relation to any freehold or leasehold property comprised in the Charged Assets;',
            'not make any structural alterations or additions to the Scheduled Property or any part of the Scheduled Property without the prior written consent of the Lender;',
            'not suffer any change of use of or carry out any works in or upon any of the Scheduled Property which would necessitate obtaining planning permission pursuant to the Planning Acts without first obtaining such permission and, where such permission is obtained, ensure that such permission is fully complied with and that documentary evidence of such compliance is furnished to the Lender;',
            'not grant or enter into any easements, wayleaves, servitudes or similar arrangements in respect of any freehold or leasehold property comprised in the Charged Assets without the prior written consent of the Lender;',
            'not do or allow to be done any act which could have a material adverse effect on the value of the Scheduled Property or as a result of which any lease of the Scheduled Property may become liable to forfeiture or otherwise be terminated;',
            'not agree to the terms of any rent review or agree to or permit any alteration, variation or addition to the terms of any Occupational Lease without the prior written consent of the Lender;',
            'promptly pay all rates, rents, taxes and other outgoings in respect of the Scheduled Property;',
            'not to agree to the compulsory purchase of the whole or any part of the Scheduled Property owned by the Chargor without the prior written consent of the Lender;'
        ]

        for i, undertaking in enumerate(property_undertakings, 1):
            cls._add_numbered_list_item(doc, f'7.3.{i}', undertaking)

        # Add 7.3.19 - complex subsection
        cls._add_numbered_list_item(doc, '7.3.19',
                                    'observe and perform all covenants and stipulations from time to time affecting the Scheduled Property and not without the prior consent in writing of the Lender:')

        property_19_subitems = [
            'enter into any onerous or restrictive obligations affecting the Scheduled Property;',
            'create or permit to arise any overriding interest or any easement or right whatsoever in or on the Scheduled Property which might adversely affect the value of the Scheduled Property;',
            'do or suffer to be done on the Scheduled Property anything which is a Development;',
            'do or suffer or omit to be done any act, matter or thing whereby any provision of law from time to time in force affecting the Scheduled Property is infringed; nor',
            'do or suffer to be done any act or thing whereby any fee farm grant or lease is likely to become liable to forfeiture or otherwise be unilaterally determined by the landlord or fee farm grantor as the case may be;'
        ]

        for i, subitem in enumerate(property_19_subitems, 1):
            cls._add_numbered_sublist_item(doc, chr(96 + i), subitem)  # a, b, c, d, e

        # Continue with remaining property undertakings
        remaining_property_undertakings = [
            'keep all buildings, structures, fixtures and fittings (including fixtures and fittings) In good and substantial repair and in good working order and condition (fair wear and tear excepted) and not pull down or remove or materially alter or sell or otherwise dispose of any of the same without the prior consent in writing of the Lender except in the ordinary course of use, repair, maintenance or improvement and not do or omit to do anything which could reasonably be expected to result in any item of or any part thereof being confiscated, seized, requisitioned, taken in execution, impounded or otherwise taken out of the Chargor\'s control;',
            'notify the Lender of any notice received in relation to the Scheduled Property which might reasonably be expected to adversely affect the value of the Scheduled Property or the Security in the Scheduled Property and, within thirty days after receipt by the Chargor of any application, requirement, order or notice served or given by any public or local or any other authority with respect to any asset secured by this Deed (or any material part of such an asset), give written notice thereof to the Lender and also (within seven days after demand) produce such notice or a copy of such notice to the Lender and inform it of the steps taken or proposed to be taken to comply with any such requirement thereby made or implicit in such notice;',
            'If all or any of the Scheduled Property is subject to an application for first registration in the Property Registration Authority, to progress such application as expeditiously as possible and keep the Lender informed of the progress of such application.'
        ]

        for i, undertaking in enumerate(remaining_property_undertakings, 20):
            cls._add_numbered_list_item(doc, f'7.3.{i}', undertaking)

            # 7.4 Prescribed Form Charge
        cls._add_subsection_header(doc, '7.4', 'Prescribed Form Charge')
        cls._add_subsection_content(doc,
                                    'The Chargor irrevocably covenants and undertakes with the Lender that it shall at any time, if called upon to do so by the Lender, execute over all or any part of any property acquired by it after the date of this Deed a legal charge in favour of the Lender in such form as the Lender may require (including by way of a charge in the Prescribed Form);')

    @classmethod
    def _add_rights_of_enforcement_section(cls, doc, context):
        """Add Clause 8 - RIGHTS OF ENFORCEMENT"""
        cls._add_main_clause_header(doc, '8', 'RIGHTS OF ENFORCEMENT')

        cls._add_subsection_content(doc,
                                    'The Security will become immediately enforceable upon the occurrence of an Event of Default which shall include any of the following events, each of which shall constitute an Event of Default:')

        # Events of Default
        events_of_default = [
            ('Failure to pay',
             'if the Chargor fails to pay all or any of the Secured Liabilities when due unless such failure to pay occurs solely as a result of any technical or administrative difficulties relating to the transfer of money and such default is remedied within two (2) Business Days of its due date;'),
            ('Breach of representation or warranty',
             'if any representation, warranty, statement, information or certificate made or provided by the Chargor pursuant to this Deed and/or any other Finance Document to which it is a party or any other document delivered by or on behalf of the Chargor in connection with any Deed and/or any other Finance Document proves to have been incorrect or misleading in any respect when made or deemed to be repeated;'),
            ('Breach of covenant',
             'if the Chargor breaches or fails to observe or perform any of the covenants in this Deed and/or any other Finance Document to which it is a party;'),
            ('Repudiation',
             'if the Chargor repudiates this Deed and/or any Finance Document to which it is a party or any agreement or document contemplated by this Deed and/or any Finance Document or does or causes to be done any act or thing evidencing an intention to repudiate any such document; or'),
            ('Insolvency', '''if:
    
       (a) the Chargor is unable or admits inability to pay its debts as they fall due or otherwise becomes insolvent;
    
       (b) the Chargor stops or suspends or threatens to stop or suspend payment of its debts;
    
       (c) a moratorium is declared in respect of any of the Chargor's Financial Indebtedness; or
    
       (d) the Chargor, as a result of anticipated financial difficulties, begins negotiations with any creditor for the rescheduling of any of its Financial Indebtedness.'''),
            ('Insolvency proceedings', '''if:
    
       (a) the Chargor is declared a bankrupt or has entered into, initiated or been the subject of any insolvency arrangements within the meaning of the Personal Insolvency Act, 2012;
    
       (b) any formal or legal step or other procedure is taken with a view to moratorium or a composition, compromise, assignment or similar arrangement with any creditors of the Chargor;
    
       (c) a receiver is appointed in respect of the assets of the Chargor; or
    
       (d) any other analogous step or procedure is taken in any jurisdiction;'''),
            ('Prior charge holders',
             'if the holder of any Security Interest from or over any part of the property, assets or undertaking of the Chargor takes possession of, or formally indicates an intention to take possession of, any part of the property, assets or undertaking of the Chargor and any charge, whether fixed or floating, granted by the Chargor over its property, assets and undertaking, crystallises or becomes enforceable or if any other action is taken to enforce any Security Interest granted, created or issued by the Chargor;'),
            ('Material adverse change',
             'if there occurs or is likely to occur a material adverse change in the business, operations, financial condition or financial prospects of the Chargor which the Lender reasonably believes is likely to affect the Chargor\'s ability to meet its obligations under this Deed and/or any Finance Document to which it is a party;'),
            ('Illegality/validity',
             'if this Deed and/or any Finance Document to which it is a party fails or ceases in any respect to have full force and effect or to be continuing or is terminated or is disputed or becomes jeopardised, invalid or unenforceable;'),
            ('Litigation',
             'if any litigation, arbitration or administrative proceedings or any dispute affecting the Chargor or any of its assets, rights or revenues are commenced or threatened which, if determined against the Chargor, might be expected to materially and adversely affect its business, operations, financial conditions or financial prospects or its ability to duly perform its obligations under this Deed and/or any Finance Document to which it is a party;'),
            ('Seizure of property',
             'if any distress, execution, attachment or other analogous legal process is levied, enforced or sued upon or against any part of the property, assets or undertaking of the Chargor;'),
            ('Default in payment of taxes',
             'if the Chargor defaults in the payment of any taxes due and payable (other than those being contested bona fide and in good faith and where the Chargor has made adequate reserves therefor).')
        ]

        for i, (event_name, event_desc) in enumerate(events_of_default, 1):
            cls._add_numbered_event_of_default(doc, str(i), event_name, event_desc)

    @classmethod
    def _add_enforcement_of_security_section(cls, doc, context):
        """Add Clause 9 - ENFORCEMENT OF SECURITY"""
        cls._add_main_clause_header(doc, '9', 'ENFORCEMENT OF SECURITY')

        # 9.1 Powers of Lender
        cls._add_subsection_header(doc, '9.1', 'Powers of Lender')
        cls._add_subsection_content(doc,
                                    'After the Security has become enforceable in accordance with clause 8 (Rights of Enforcement), the Lender may in its absolute discretion and without prior notice to the Chargor, enforce and realise all or any part of the Security and/or take possession of, hold or dispose of all or any of the Charged Assets in any manner it sees fit and the statutory powers of sale and of appointing a Receiver and other powers conferred on mortgagees by the Act shall apply to this Deed in each case as varied by this Deed. Section 99 of the Act shall not apply to this Deed and neither the Lender nor any Receiver shall be obliged to take any steps to sell or lease the Charged Assets after going into possession of same and the Lender and any Receiver shall have absolute discretion as to the time of exercise of the power of sale and the power of leasing and all other powers conferred on them by the Act or otherwise.')

        # 9.2 Additional Power of the Lender
        cls._add_subsection_header(doc, '9.2', 'Additional Power of the Lender')

        cls._add_numbered_subsection(doc, '9.2.1',
                                     'All or any of the powers, authorities and discretions which are conferred by this Deed, either expressly or impliedly, upon a Receiver, may be exercised by the Lender in relation to the whole of the Charged Assets or any part thereof without first appointing a Receiver of such Charged Assets or any part thereof or notwithstanding the appointment of a Receiver of such Charged Assets or any part thereof.')

        cls._add_numbered_subsection(doc, '9.2.2',
                                     'The powers conferred by this Deed in relation to the Charged Assets or any part thereof on the Lender or on any Receiver of such Charged Assets or any part thereof shall be in addition to, and not in substitution for, the powers conferred on mortgagees or receivers by any law.')

        # 9.3 Lender as mortgagee in possession
        cls._add_subsection_header(doc, '9.3', 'Lender as mortgagee in possession')
        cls._add_subsection_content(doc,
                                    'At any time after the security hereby constituted has become enforceable in accordance with clause 8 (Rights of Enforcement) and without the need to obtain the consent of the Chargor or an order for possession under section 97 or 98 of the Act, the Lender may without further notice or demand enter into possession of the Charged Assets. The rights of the Lender under this clause are without prejudice to, and/or in addition to, any right of possession (express or implied) to which it is at any time otherwise entitled (whether by virtue of this Deed, operation of law, statute, contract or otherwise) to enter into possession of the Charged Assets or any part of the Charged Assets and the Lender shall have power to:')

        lender_powers = [
            'enter upon or take possession of and hold any of the Charged Assets or any part of the Charged Assets and carry out any such repairs, amendments, alterations and additions as the Lender shall reasonably consider necessary or desirable for the maintenance or protection of the same or any part of the Charged Assets;',
            'demise or agree to demise any of the Charged Assets or any part of the Charged Assets of which the Lender is in possession for such period at such rent and upon such terms with or without a premium or fine in all respects as the Lender may from time to time think fit;',
            'carry on the business of the Chargor and manage and conduct the same as it shall in its sole discretion think fit; and',
            'do all such other acts and things which, in the opinion of the Lender, are incidental to any of the powers, functions, authorities or discretions conferred on the Lender pursuant to this Deed or by statute and law generally.'
        ]

        for i, power in enumerate(lender_powers, 1):
            cls._add_numbered_list_item(doc, str(i), power)

        # 9.4 Power of Sale
        cls._add_subsection_header(doc, '9.4', 'Power of Sale')
        cls._add_subsection_content(doc,
                                    'At any time after the security hereby constituted has become enforceable in accordance with clause 8 (Rights of Enforcement) the power of sale and all other powers conferred on mortgagees by the Act shall be exercisable immediately without the need:')

        sale_powers = [
            'for the occurrence of any of the events specified in paragraphs (a) to (c) of section 100(1) of the Act;',
            'to give notice as specified in the final proviso to section 100(1) of the Act;',
            'to obtain the consent of the Chargor or a court order authorising the exercise of the power of sale under sections 100(2) or (3) of the Act; or',
            'to give any notice to the Chargor under section 103(2) of the Act.'
        ]

        for i, power in enumerate(sale_powers, 1):
            cls._add_numbered_list_item(doc, str(i), power)

        cls._add_subsection_content(doc, 'Sections 93, 94 and 95 of the Act shall not apply to this Deed.')

        # 9.5 Power of Leasing and accepting Surrenders
        cls._add_subsection_header(doc, '9.5', 'Power of Leasing and accepting Surrenders')
        cls._add_subsection_content(doc,
                                    'The statutory powers of leasing conferred on the Lender and any Receiver are extended so as to authorise the Lender and any Receiver to lease, make arrangements for leases, accept surrenders of leases and make agreements to accept surrenders of leases as it or he may think fit and without the need to comply with any provision of sections 112 to 114 of the Act. Without prejudice to the generality of the foregoing the Lender and any Receiver may exercise the statutory power to accept surrenders of leases conferred by the Act for any purpose that it or he thinks fit and not just for the purpose of granting new leases under section 112 of the Act and any new lease granted by the Lender or any Receiver following the acceptance of a surrender need not comply with the requirements of section 114(3) of the Act.')

        # 9.6 Liability and Privileges of the Lender and Receiver
        cls._add_subsection_header(doc, '9.6', 'Liability and Privileges of the Lender and Receiver')

        cls._add_numbered_subsection(doc, '9.6.1',
                                     'Nothing in this Deed shall be deemed to impose on the Lender or any Receiver, any liability whatsoever in relation to the Charged Assets or render the Lender or any Receiver liable to account to the Chargor as mortgagee in possession in respect of any Charged Assets or be liable to the Chargor in respect of any loss or damage which arises out of the exercise, the attempted or purported exercise or the failure to exercise any of their respective powers or for any other loss of any nature whatsoever.')

        cls._add_numbered_subsection(doc, '9.6.2',
                                     'The Lender will not be liable for any involuntary losses which may occur as a result of the exercise or execution of the statutory power of sale or any of the powers or trust expressed or implied which may be vested in the Lender by virtue of this Deed.')

        cls._add_numbered_subsection(doc, '9.6.3',
                                     'The Lender and any Receiver appointed under this Deed shall be entitled to all the rights, powers, privileges and immunities conferred by the Act on mortgagees and receivers when such receivers have been duly appointed under the Act, but so that the power of sale and other powers conferred by the Act and as may be available at law shall be as varied and extended by this Deed.')

        # 9.7 Protection of third parties
        cls._add_subsection_header(doc, '9.7', 'Protection of third parties')

        cls._add_numbered_subsection(doc, '9.7.1',
                                     'No purchaser or other person will be bound or concerned to see or enquire whether the right of the Lender or any Receiver appointed by the Lender to exercise any of the powers conferred by this Deed has arisen or not or be concerned with notice to the contrary or with the propriety of the exercise or purported exercise of such powers.')

        cls._add_numbered_subsection(doc, '9.7.2',
                                     'The receipt of the Lender or any Receiver shall be an absolute and complete discharge to a purchaser and shall relieve it of any obligation to see to the application of any monies paid to or at the direction of the Lender or any Receiver.')

        cls._add_numbered_subsection(doc, '9.7.3',
                                     'All protections to purchasers contained in sections 105, 106 and 108(5) of the Act shall apply to any person (including a purchaser) dealing with the Lender or any Receiver in like manner as if the statutory powers of sale and appointing a Receiver had not been varied or extended by this Deed.')

        # 9.8 Delegation
        cls._add_subsection_header(doc, '9.8', 'Delegation')

        cls._add_numbered_subsection(doc, '9.8.1',
                                     'The Lender or any Receiver may from time to time delegate by power of attorney or otherwise to any person or corporation any of the powers and discretions of the Lender or any Receiver under this Deed whether arising by statute or otherwise upon such terms and for such periods of time as it may think fit and may determine by such delegation.')

        cls._add_numbered_subsection(doc, '9.8.2',
                                     'Neither the Lender nor any Receiver will be liable to the Chargor for any loss or damage arising from any act, default, omission or misconduct of any such delegate and references in this Deed to the Lender or to any Receiver will, where the context so admits, include reference to any delegates so appointed.')

    @classmethod
    def _add_receivers_section(cls, doc, context):
        """Add Clause 10 - RECEIVERS"""
        cls._add_main_clause_header(doc, '10', 'RECEIVERS')

        # 10.1 Appointment and removal
        cls._add_subsection_header(doc, '10.1', 'Appointment and removal')
        cls._add_subsection_content(doc,
                                    'At any time after the Security has become enforceable in accordance with clause 8 (Rights of Enforcement), the Lender (without the need for the occurrence of any of the events specified in paragraphs (a) to (c) of section 108(1) of the Act) may by instrument in writing (under seal or otherwise under the hand of any officer, manager or other nominated person of the Lender), without prior notice to the Chargor:')

        appointment_powers = [
            'appoint one or more persons considered by it to be competent to be a Receiver (which shall, for the avoidance of doubt include a receiver and manager or joint receiver) of the whole or any part of the Charged Assets; and',
            'remove any Receiver so appointed and appoint another or others in his place or appoint another or others to act jointly with such Receiver provided that where more than one Receiver is appointed they shall have the power to act severally unless the Lender shall otherwise specify.'
        ]

        for i, power in enumerate(appointment_powers, 1):
            cls._add_numbered_list_item(doc, str(i), power)

        cls._add_subsection_content(doc,
                                    'The foregoing powers of appointment of a Receiver shall be in addition to and not to the prejudice of all statutory and other powers of the Lender under the Act or otherwise and such powers as varied or extended by this Deed shall be and remain exercisable by the Lender in respect of any of the Charged Assets notwithstanding the appointment of a Receiver over any of the Charged Assets.')

        # 10.2 Receiver as agent of the Chargor
        cls._add_subsection_header(doc, '10.2', 'Receiver as agent of the Chargor')
        cls._add_subsection_content(doc,
                                    'Any Receiver appointed under this Deed shall at all times be the agent of the Chargor and the Chargor shall be solely responsible for his acts and defaults and liable on any contract or engagements made or entered into or adopted by him and the Receiver shall at no time act as agent for the Lender.')

        # 10.3 Remuneration of the Receiver
        cls._add_subsection_header(doc, '10.3', 'Remuneration of the Receiver')
        cls._add_subsection_content(doc,
                                    'The Lender may fix the remuneration of any Receiver appointed by it and direct payment of the Receiver out of the Charged Assets or any part of the Charged Assets, but the Chargor alone will be liable for the payment of such remuneration and the provisions of section 108(7) of the Act shall not apply to this Deed.')

        # 10.4 Powers of the Receiver
        cls._add_subsection_header(doc, '10.4', 'Powers of the Receiver')
        cls._add_subsection_content(doc,
                                    'A Receiver shall, in relation to the Charged Assets over which he is appointed, have all powers conferred by the Act and all other statutes in the same way as if the Receiver had been duly appointed under the Act and shall be entitled to exercise such powers in such manner and on such terms as he may in his absolute discretion think fit. Furthermore a Receiver shall have (in each case at the cost of the Chargor) the following additional powers:')

        receiver_powers = [
            'to enter on, take possession of, collect and get in all or any part of the property in respect of which the Receiver is appointed and for that purpose take any proceedings in the name of the Chargor or otherwise as may seem expedient;',
            'to carry on or manage or develop or diversify or concur in carrying on or managing or developing or diversifying the business of the Chargor in respect of all or any part of the property over which the Receiver is appointed and for that purpose raise money on any part of the property in respect of which the Receiver is appointed in priority to this Security or otherwise;',
            'to raise and borrow money for any other purpose, whether secured on the security of any of the Charged Assets or not and either in priority to the security constituted by this Deed or otherwise and generally on any terms and for whatever purpose consistent with his appointment which he thinks fit;',
            'to appoint, hire and employ officers, employees, contractors, agents and advisors of all kinds as the Receiver shall deem necessary or appropriate and to discharge any such persons and any such persons appointed, hired or employed by the Chargor;',
            'to enter on or otherwise take possession of the Charged Assets to make and effect any repairs, renewals, improvements, add to or develop or to complete any Work-in-Progress or building or structure which may be unfinished and to maintain or renew all Insurances;',
            'to employ, hire and appoint officers, employees, contractors, agents and advisors to assist in carrying on and managing the business of the Chargor and to terminate any appointment or contract of employment (whether or not pre-dating his appointment as Receiver);',
            'to redeem any Security Interest on, over or affecting the Charged Assets or any part of the Charged Assets;',
            '''to promote or procure the formation of any new company and, in the case of such new company:
    
    (a) to subscribe for or acquire (for cash or otherwise) any investment in such new company;
    
    (b) to sell, transfer, assign, hire out and lend, and grant leases, tenancies and rights of user of, the Charged Assets to any such new company and accept as consideration or part of the consideration therefor any shares or other security in or of any company or allow the payment of the whole or any part of such consideration to remain deferred or outstanding; and
    
    (c) to sell, transfer, assign, exchange and otherwise dispose of or realise any such shares or other security or deferred consideration or part of such shares or other security or deferred consideration or any rights attaching to such shares or other security or deferred consideration;''',
            'to grant any lease, licence or tenancy or right of or affecting the Charged Assets for any term or term of years at any or no rent or fee and with or without any premium and accept the surrender of any lease or tenancy or right and give a valid receipt for any premium payable on such grant or surrender and to amend or vary any lease, licence, agreement or other arrangement in any way relating to or affecting the Charged Assets;',
            'to sell (whether by public auction, private contract or otherwise) all or any of the Charged Assets on any terms and for any consideration (including for deferred consideration or a consideration payable wholly or partly in instalments or consisting in whole or in part of shares or securities of any other company or of any other non-cash asset) on such terms and conditions as he may think fit (including conditions excluding or restricting the personal liability of the Receiver or the Lender);',
            'to enter and perform, repudiate, rescind or vary such contracts and arrangements to which the Chargor is a party or incur any obligations in relation to such contracts and/or arrangements;',
            'take any indemnity from the Chargor from and against all actions, claims, expenses, demands and liabilities whether arising out of contract or out of tort or in any other way incurred by the Receiver or by any manager, agent, officer, servant or workman for whose debt, default or miscarriage he may be answerable for anything done or omitted to be done in the exercise or purported exercise of his powers under this Deed or under any appointment duly made by the Receiver and if he thinks fit but without prejudice to the foregoing to effect with any insurance company or office or underwriters any policy or policies of insurance either in lieu or satisfaction of or in addition to such indemnity from the Chargor;',
            'to disclaim, abandon or disregard all or any of the outstanding contracts of the Chargor and to allow time for payment by or to the Chargor of any debts either with or without security;',
            'to settle, adjust, refer to arbitration, compromise and arrange any claims, accounts, disputes, questions and demands with or by any person who is or claims to be a creditor of the Chargor or relating in any way to the Charged Assets or otherwise as the Lender or the Receiver may think expedient;',
            'to bring, prosecute, enforce, defend and abandon actions, suits and proceedings in relation to the Charged Assets (or any part of the Charged Assets) or any of the businesses of the Chargor;',
            'to exercise, in relation to any Charged Assets, all the powers, authorities and things which he would be capable of exercising if he were the absolute beneficial owner of that Charged Asset; and',
            'to do all such other acts and things which, in the opinion of the Receiver, are incidental to any of the powers, functions, authorities or discretions conferred on or vested in the Receiver pursuant to this Deed or upon receivers by statute or law generally including the bringing or defending of proceedings in the name of, or on behalf of, the Chargor and the preservation, improvement, collection and/or realisation of Charged Assets and the execution of documents in the name of the Chargor (whether by hand or under seal of the Chargor).'
        ]

        for i, power in enumerate(receiver_powers, 1):
            cls._add_numbered_list_item(doc, str(i), power)

        # 10.5 Application of Monies by Receiver
        cls._add_subsection_header(doc, '10.5', 'Application of Monies by Receiver')
        cls._add_subsection_content(doc,
                                    'All monies received by any Receiver shall, notwithstanding section 109 of the Act, be applied by him for the following purposes (subject to the claims of secured and unsecured creditors (if any) ranking in priority to or pari passu with the security hereby constituted) in the following order:')

        application_order = [
            'in payment of all costs, charges and expenses of and incidental to the appointment of any Receiver and the exercise of all or any of the powers aforesaid and of all outgoings paid by any Receiver;',
            'in payment of remuneration to any Receiver at such rate as may be agreed between him and the Lender;',
            'in or towards payment and discharge of the Secured Liabilities; and',
            'in payment of any surplus to the Chargor or any other person lawfully entitled to such payment.'
        ]

        for i, item in enumerate(application_order, 1):
            cls._add_numbered_list_item(doc, str(i), item)

    @classmethod
    def _add_set_off_section(cls, doc, context):
        """Add Clause 11 - SET-OFF"""
        cls._add_main_clause_header(doc, '11', 'SET-OFF')
        cls._add_subsection_content(doc,
                                    'Without prejudice to any right of set-off or any similar right to which the Lender may be entitled at law or in equity and without prejudice to anything else in this Deed, the Lender may at any time after the Security has become enforceable in accordance with clause 8 (Rights of Enforcement) and without further notice to or further authorisation from the Chargor, combine and consolidate all or any accounts of the Chargor with the Lender and/or set-off any monies in such accounts against any monies owed by the Chargor (whether actual or contingent) to the Lender, regardless of the place of payment or currency of either obligation. If the obligations are in different currencies, the Lender may convert either obligation at a market rate of exchange in its usual course of business for the purpose of the set-off.')

    @classmethod
    def _add_release_of_security_section(cls, doc, context):
        """Add Clause 12 - RELEASE OF SECURITY"""
        cls._add_main_clause_header(doc, '12', 'RELEASE OF SECURITY')

        # 12.1 Release of Security
        cls._add_subsection_header(doc, '12.1', 'Release of Security')
        cls._add_subsection_content(doc,
                                    'Subject to and without prejudice to clause 12.3 (Avoidance of payments), after the Secured Liabilities have been unconditionally and irrevocably paid and discharged in full and all the commitments of the Lender cancelled, the Lender shall, as soon as reasonably practicable thereafter and at the request and cost of the Chargor, execute and do all such deeds, acts and things as may be necessary to release the Security.')

        # 12.2 Entitlement to retain Security
        cls._add_subsection_header(doc, '12.2', 'Entitlement to retain Security')
        cls._add_subsection_content(doc,
                                    'If any payment or discharge of the Secured Liabilities is, in the reasonable opinion of the Lender, liable to be avoided or invalidated under any enactment relating to bankruptcy or insolvency, the Lender may refuse to grant any release of the Security for such further period as the risk of such avoidance or invalidity continues.')

        # 12.3 Avoidance of payments
        cls._add_subsection_header(doc, '12.3', 'Avoidance of payments')
        cls._add_subsection_content(doc,
                                    'No assurance, security or payment which may be avoided or adjusted under law, including under any statute relating to bankruptcy or insolvency and no release, settlement or discharge given or made by the Lender on the faith of any such assurance, security or payment, shall prejudice or affect the right of the Lender to recover the Secured Liabilities from the Chargor (including any monies which the Lender may be compelled to pay or refund under the laws of insolvency and any costs payable by it pursuant to or otherwise incurred in connection therewith) or to appoint a Receiver and enforce the Security to the full extent of the Secured Liabilities.')

    @classmethod
    def _add_waiver_of_defences_section(cls, doc, context):
        """Add Clause 13 - WAIVER OF DEFENCES"""
        cls._add_main_clause_header(doc, '13', 'WAIVER OF DEFENCES')

        cls._add_subsection_content(doc,
                                    'The obligations of the Chargor under this Deed will not be affected by any act, omission or circumstances which, but for this clause 13 (Waiver of Defences), might operate to release or otherwise exonerate the Chargor from its obligations under this Deed or affect such obligations in whole or in part including and whether or not known to the Chargor or the Lender:')

        waiver_items = [
            'any time, waiver or consent granted to or composition with any other person;',
            'the release of any persons from their obligations under any Finance Document or any Security Interest or guarantee granted in connection therewith;',
            'the taking, variation, compromise, exchange, renewal or release of, or refusal or neglect to perfect, take up or enforce, any rights against or security over assets of any person or any non-presentation or non-observance of any formality or other requirement in respect of any instrument or any failure to realise the full value of any security;',
            'any incapacity or lack of power, authority or legal personality of or dissolution or change in the members or status of any person;',
            'any amendment, novation, supplement, extension, restatement (however fundamental and whether or not more onerous) or replacement of any Finance Document or any other document or security instrument including any change in the purpose of, any extension of or any increase in any facility or the addition of any new facility under a Finance Document or other document or security;',
            'any unenforceability, illegality or invalidity of any obligation of any person under a Finance Document or any other document or security instrument;',
            'any insolvency or similar proceedings; or',
            'any other act, event or omission which, but for this clause 13 (Waiver of Defences) might operate to discharge, impair or otherwise affect any of the obligations of the Chargor under this Deed or any of the rights, powers or remedies conferred upon the Lender by a Finance Document or by law.'
        ]

        for i, item in enumerate(waiver_items, 1):
            cls._add_numbered_list_item(doc, str(i), item)

    @classmethod
    def _add_new_account_section(cls, doc, context):
        """Add Clause 14 - NEW ACCOUNT"""
        cls._add_main_clause_header(doc, '14', 'NEW ACCOUNT')
        cls._add_subsection_content(doc,
                                    'If the Lender receives, or is deemed to be affected by actual or constructive notice of any subsequent Security Interest or other interest affecting the Charged Assets (or any part of the Charged Assets), the Lender may open a new account for the Chargor. If the Lender does not open a new account then, unless the Lender gives express written notice to the contrary to the Chargor, the Lender will be treated as if it had done so at the time when it received or was deemed to have received notice and as from that time all monies paid by the Chargor shall be credited or be treated as having been credited to the new account and will not operate to reduce the amount due from the Chargor to the Lender at the time when the Lender received or was deemed to have received that notice and furthermore the Lender may forthwith discontinue any guarantee or any other facility given or granted on the account of the Chargor.')

    @classmethod
    def _add_application_of_proceeds_section(cls, doc, context):
        """Add Clause 15 - APPLICATION OF PROCEEDS"""
        cls._add_main_clause_header(doc, '15', 'APPLICATION OF PROCEEDS')

        # 15.1 Application of proceeds
        cls._add_subsection_header(doc, '15.1', 'Application of proceeds')
        cls._add_subsection_content(doc,
                                    'Any monies received by the Lender after the Security has become enforceable in accordance with clause 8 (Rights of Enforcement), shall notwithstanding the provisions of section 107 of the Act be applied in the following order of priority:')

        application_items = [
            'in payment of all costs (including break-costs), charges and expenses incurred by the Lender or any Receiver under or in connection with this Deed including all remuneration due to any Receiver;',
            'in or towards payment of the Secured Liabilities; and',
            'in payment of the surplus (if any) to the Chargor or any other person entitled to it.'
        ]

        for i, item in enumerate(application_items, 1):
            cls._add_numbered_list_item(doc, str(i), item)

        # 15.2 Other claims
        cls._add_subsection_header(doc, '15.2', 'Other claims')
        cls._add_subsection_content(doc,
                                    'Without prejudice to the right of the Lender to recover any shortfall from the Chargor, the provisions of clause 15.1 (Application of proceeds) are subject to the payment of any claims having priority over the Security.')

    @classmethod
    def _add_suspense_account_section(cls, doc, context):
        """Add Clause 16 - SUSPENSE ACCOUNT"""
        cls._add_main_clause_header(doc, '16', 'SUSPENSE ACCOUNT')
        cls._add_subsection_content(doc,
                                    'Any monies received, recovered or realised by the Lender under this Deed (including the proceeds of any conversion of currency) may, at the discretion of the Lender, be credited to any interest-bearing suspense account maintained with any bank, building society or financial institution as it considers appropriate and may be held in such account for so long as the Lender may think fit pending the application of such monies from time to time (as the Lender is entitled to do in its discretion) in or towards the discharge of the Secured Liabilities (or any part of the Secured Liabilities) and the Chargor shall not be entitled to withdraw any amount at any time standing to the credit of any suspense account referred to above.')

    @classmethod
    def _add_power_of_attorney_section(cls, doc, context):
        """Add Clause 17 - POWER OF ATTORNEY"""
        cls._add_main_clause_header(doc, '17', 'POWER OF ATTORNEY')

        # 17.1 Appointment and powers
        cls._add_subsection_header(doc, '17.1', 'Appointment and powers')
        cls._add_subsection_content(doc,
                                    'The Chargor by way of security hereby irrevocably appoints (in the case of those matters referred to in clause 17.1.2 with immediate effect but otherwise from the Security becoming enforceable in accordance with clause 8 (Rights of Enforcement) the Lender and every Receiver jointly and severally to be its attorney in its name and on its behalf:')

        attorney_powers = [
            'to execute and complete any documents or instruments to which the Lender or such Receiver may require for perfecting the title of the Lender to the Charged Assets or for vesting the same in the Lender, its nominees or any purchasers;',
            'to sign, execute, seal and deliver and otherwise perfect any further security document or instrument required to be provided to the Lender pursuant to clause 5.2 (Further assurances); and',
            'otherwise generally to sign, seal, execute and deliver all deeds, assurances, agreements and documents and to do all acts and things which may be required for the full exercise of all or any of the powers conferred on the Lender or a Receiver under this Deed or which may be deemed expedient by the Lender or a Receiver in connection with any disposition, realisation or getting in by the Lender or such Receiver of such Charged Assets or any part of such Charged Assets or in connection with any other exercise of any power under this Deed.'
        ]

        for i, power in enumerate(attorney_powers, 1):
            cls._add_numbered_list_item(doc, str(i), power)

        # 17.2 Ratification
        cls._add_subsection_header(doc, '17.2', 'Ratification')
        cls._add_subsection_content(doc,
                                    'The Chargor ratifies and confirms whatever any such attorney appointed under clause 17.1 (Appointment and powers) shall do or purport to do in the exercise or purported exercise of all or any of the powers, authorities and discretions referred to in such clause.')

    @classmethod
    def _add_expenses_and_indemnity_section(cls, doc, context):
        """Add Clause 18 - EXPENSES AND INDEMNITY"""
        cls._add_main_clause_header(doc, '18', 'EXPENSES AND INDEMNITY')

        # 18.1 Indemnity
        cls._add_subsection_header(doc, '18.1', 'Indemnity')
        cls._add_subsection_content(doc,
                                    'The Chargor shall indemnify the Lender and any Receiver appointed by it against all actions, claims, demands, losses, expenses or liabilities of whatever nature after the date of this Deed incurred by them or any officer, agent or employee for whose liability act or omission they or any of them may be answerable for anything done or omitted in:')

        indemnity_items = [
            'the purported exercise of the powers contained in this Deed;',
            'occasioned by any breach by the Chargor of any of its covenants or other obligations to the Lender;',
            'any failure on the part of the Lender to perform any obligations of the Chargor as a result of the grant of Security under this Deed; and',
            'in consequence of any payment in respect of the Secured Liabilities (whether made by the Chargor or a third person) being declared void or impeached for any reason unless such loss or damage shall be caused by the Lender\'s or Receiver\'s own fraud or wilful neglect or gross negligence.'
        ]

        for i, item in enumerate(indemnity_items, 1):
            cls._add_numbered_list_item(doc, str(i), item)

    @classmethod
    def _add_currencies_section(cls, doc, context):
        """Add Clause 19 - CURRENCIES"""
        cls._add_main_clause_header(doc, '19', 'CURRENCIES')

        cls._add_numbered_subsection(doc, '19.1',
                                     'All monies received or held by the Lender or a Receiver under this Deed may from time to time after demand has been made be converted into such other currency as the Lender considers necessary or desirable to cover the obligations of the Chargor in that currency at the then prevailing spot rate of exchange (as conclusively determined by the Lender) for purchasing the currency to be acquired with the existing currency. The Chargor shall indemnify the Lender against the full euro cost (including all costs charges and expenses) incurred in relation to such conversion of currency. Neither the Lender nor any Receiver shall have any liability to the Chargor in respect of any loss resulting from any fluctuation in exchange rates after such sale.')

        cls._add_numbered_subsection(doc, '19.2',
                                     'No payment to the Lender (whether under any judgment or court order or otherwise) will discharge the obligation or liability of the Chargor in respect of which it was made and until the Lender receives payment in full in the currency in which such obligation or liability was originally incurred and to the extent that the amount of any such payment, on actual conversion into such currency, falls short of such obligation or liability expressed in that currency, the Lender shall have a further separate cause of action against the Chargor and will be entitled to enforce the charges created by this Deed to recover the amount of any shortfall.')

    @classmethod
    def _add_transfers_section(cls, doc, context):
        """Add Clause 20 - TRANSFERS"""
        cls._add_main_clause_header(doc, '20', 'TRANSFERS')

        # 20.1 Transfer by Chargor
        cls._add_subsection_header(doc, '20.1', 'Transfer by Chargor')
        cls._add_subsection_content(doc,
                                    'The Chargor may not assign or otherwise transfer all or any of its rights, benefits or obligations under this Deed.')

        # 20.2 Transfer by Lender
        cls._add_subsection_header(doc, '20.2', 'Transfer by Lender')

        cls._add_numbered_subsection(doc, '20.2.1',
                                     'The Lender may at any time assign and transfer all or any of its rights or obligations under this Deed to any person and shall be entitled to disclose such information concerning the Chargor under this Deed as the Lender considers appropriate to any such person, including any actual or prospective direct or indirect successor, assignee or transferee or any person to whom the information may be required to be disclosed by any applicable law.')

        cls._add_numbered_subsection(doc, '20.2.2',
                                     'In the event of assignment or transfer by the Lender pursuant to clause 20.2.1, the Chargor shall at the request of the Lender join in such assignment, transfer or other document so as to cause full beneficial title to the Security to be passed to the relevant assignee or transferee.')

    @classmethod
    def _add_notices_section(cls, doc, context):
        """Add Clause 21 - NOTICES"""
        cls._add_main_clause_header(doc, '21', 'NOTICES')

        # 21.1 Communications in writing
        cls._add_subsection_header(doc, '21.1', 'Communications in writing')
        cls._add_subsection_content(doc,
                                    'Any communication to be made under or in connection with this Deed shall be made in writing, in the English language and may be delivered by hand, sent by prepaid post (including via airmail to another jurisdiction) or sent by email;')

        # 21.2 Addresses
        cls._add_subsection_header(doc, '21.2', 'Addresses')
        cls._add_subsection_content(doc,
                                    'The address and email address (and the department or officer, if any, for whose attention the communication is to be made) for each party for any communication or document to be made or delivered under or in connection with this Deed is that identified with its name below or any substitute address, email address or department or officer as either party may notify the other by not less than five (5) Business Days\' notice in writing:')

        # Add address table
        cls._add_address_table(doc, context)

        # 21.3 Delivery
        cls._add_subsection_header(doc, '21.3', 'Delivery')
        cls._add_subsection_content(doc,
                                    'Any communication or document made or delivered by one person to another under or in connection with this Deed will only be effective:')

        delivery_items = [
            'if delivered by hand, on delivery;',
            'if sent by prepaid post, twenty-four (24) hours after posting; or',
            'if sent by email, at the time of sending the email (provided that a return receipt is received by the person sending the email)'
        ]

        for i, item in enumerate(delivery_items, 1):
            cls._add_numbered_list_item(doc, str(i), item)

        cls._add_subsection_content(doc,
                                    'provided that if, in accordance with the above provisions, any such notice or other communication would otherwise be deemed to be given or made outside working hours (between 9am and 5pm Dublin time on a Business Day) such notice or other communication shall be deemed to be given or made at the start of working hours on the next succeeding Business Day.')

        # 21.4 Death of the Chargor
        cls._add_subsection_header(doc, '21.4', 'Death of the Chargor')
        cls._add_subsection_content(doc,
                                    'In the case of the death of the Chargor and until the Lender receives notices in writing of the grant of probate of the will or letters of administration in respect of the estate of the deceased, and in each such case, an address of communication for the Chargor\'s personal representative(s), any notice of demand by the Lender sent as aforesaid and addressed to the deceased shall for all purposes be deemed sufficient notice or demand by the Lender to the deceased and the relevant Chargor\'s personal representatives and shall be as effective as if the deceased were still living.')

    @classmethod
    def _add_miscellaneous_section(cls, doc, context):
        """Add Clause 22 - MISCELLANEOUS"""
        cls._add_main_clause_header(doc, '22', 'MISCELLANEOUS')

        # 22.1 No waivers, cumulative remedies
        cls._add_subsection_header(doc, '22.1', 'No waivers, cumulative remedies')
        cls._add_subsection_content(doc,
                                    'The rights of the Lender and any Receiver may be exercised as often as necessary, are cumulative and are in addition to its respective rights under general law. No failure or delay on the part of the Lender or any Receiver to exercise, or any partial exercise of any power, right or remedy shall operate as a waiver of that power, right or remedy or preclude its further exercise or the exercise of any other power, right or remedy;')

        # 22.2 Severability
        cls._add_subsection_header(doc, '22.2', 'Severability')
        cls._add_subsection_content(doc,
                                    'If any of the provisions of this Deed is or becomes invalid, illegal or unenforceable, that shall not affect the validity, legality or enforceability of any other provision in this Deed;')

        # 22.3 Variation
        cls._add_subsection_header(doc, '22.3', 'Variation')
        cls._add_subsection_content(doc,
                                    'This Deed may not be released, discharged, supplemented, amended, varied or modified in any matter except by an instrument in writing signed by a duly authorised officer or representative of each of the parties to this Deed;')

        # 22.4 Unfettered discretion
        cls._add_subsection_header(doc, '22.4', 'Unfettered discretion')
        cls._add_subsection_content(doc,
                                    'Save as otherwise stated in this Deed, any liability or power which may be exercised or any determination which may be made under this Deed by the Lender may be exercised or made in its absolute and unfettered discretion and it shall not be obliged to give reasons therefor.')

    @classmethod
    def _add_counterparts_section(cls, doc, context):
        """Add Clause 23 - COUNTERPARTS"""
        cls._add_main_clause_header(doc, '23', 'COUNTERPARTS')
        cls._add_subsection_content(doc,
                                    'This Deed may be executed in any number of counterparts and all those counterparts taken together shall be deemed to constitute one and the same instrument.')

    @classmethod
    def _add_governing_law_section(cls, doc, context):
        """Add Clause 24 - GOVERNING LAW"""
        cls._add_main_clause_header(doc, '24', 'GOVERNING LAW')
        cls._add_subsection_content(doc,
                                    'This Deed and all relationships created hereby in all respects will be governed by and construed in accordance with the laws of Ireland.')

    @classmethod
    def _add_jurisdiction_section(cls, doc, context):
        """Add Clause 25 - JURISDICTION"""
        cls._add_main_clause_header(doc, '25', 'JURISDICTION')

        cls._add_numbered_subsection(doc, '25.1',
                                     'The Chargor hereby agrees for the exclusive benefit of the Lender that any legal action or proceedings ("Proceedings") brought against the Chargor with respect to this Deed may be brought in the High Court in Ireland or such competent Court of Ireland as the Lender may elect and the Chargor waives any objection to the bringing of Proceedings in such courts whether on grounds of venue or on the grounds that such Proceedings have been brought in an inconvenient forum. The Chargor undertakes to enter an unconditional appearance within ten (10) Business Days after the completion of any service or process in any Proceedings. The Chargor hereby consents to the service by post of any process issued in connection with this Deed. Nothing in this Deed will affect the right to serve process in any other manner permitted by law;')

        cls._add_numbered_subsection(doc, '25.2',
                                     'Nothing contained in this Deed would limit the right of the Lender to take Proceedings against the Chargor in any other court of competent jurisdiction, nor will the taking of any such Proceedings in any one or more jurisdictions preclude the taking by the Lender of Proceedings in any other jurisdiction whether concurrently or not.')

    @classmethod
    def _add_all_schedules(cls, doc, context):
        """Add all schedules"""
        # Page break before schedules
        doc.add_page_break()

        # Add all schedules
        cls._add_schedule_1(doc, context)
        cls._add_schedule_2(doc, context)
        cls._add_schedule_3(doc, context)
        cls._add_schedule_4(doc, context)

    @classmethod
    def _add_schedule_1(cls, doc, context):
        """Add Schedule 1"""
        # Schedule 1 header
        schedule_header = doc.add_paragraph()
        schedule_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        schedule_run = schedule_header.add_run('SCHEDULE 1')
        schedule_run.bold = True
        schedule_run.font.size = Pt(14)
        schedule_run.underline = True
        schedule_header.paragraph_format.space_after = Pt(12)

        # Part I
        part1_header = doc.add_paragraph()
        part1_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part1_run = part1_header.add_run('Part I\nScheduled Property')
        part1_run.bold = True
        part1_run.font.size = Pt(12)
        part1_header.paragraph_format.space_after = Pt(12)

        property_para = doc.add_paragraph()
        property_para.add_run(
            f"ALL THAT AND THOSE the property comprised in Folio {context.get('folio_number', '_' * 20)} of the register {context.get('county', '_' * 20)}.")
        property_para.paragraph_format.space_after = Pt(24)

        # Part II
        part2_header = doc.add_paragraph()
        part2_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part2_run = part2_header.add_run('Part II\nFORM 52')
        part2_run.bold = True
        part2_run.font.size = Pt(12)
        part2_header.paragraph_format.space_after = Pt(12)

        # Add Form 52 content
        cls._add_form_52_content(doc, context)

    @classmethod
    def _add_schedule_2(cls, doc, context):
        """Add Schedule 2"""
        doc.add_page_break()

        # Schedule 2 header
        schedule_header = doc.add_paragraph()
        schedule_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        schedule_run = schedule_header.add_run('SCHEDULE 2')
        schedule_run.bold = True
        schedule_run.font.size = Pt(14)
        schedule_run.underline = True
        schedule_header.paragraph_format.space_after = Pt(12)

        # Part I - Occupational Leases
        part1_header = doc.add_paragraph()
        part1_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part1_run = part1_header.add_run('Part I\nOccupational Leases')
        part1_run.bold = True
        part1_run.font.size = Pt(12)
        part1_header.paragraph_format.space_after = Pt(12)

        leases_para = doc.add_paragraph()
        leases_para.add_run('None at the date of this Debenture')
        leases_para.paragraph_format.space_after = Pt(24)

        # Part II - Insurances
        part2_header = doc.add_paragraph()
        part2_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part2_run = part2_header.add_run('Part II\nInsurances')
        part2_run.bold = True
        part2_run.font.size = Pt(12)
        part2_header.paragraph_format.space_after = Pt(12)

        # Create insurance table
        insurance_table = doc.add_table(rows=1, cols=3)
        insurance_table.style = 'Table Grid'
        insurance_table.columns[0].width = Inches(2.5)
        insurance_table.columns[1].width = Inches(2.5)
        insurance_table.columns[2].width = Inches(2.0)

        # Header row
        header_cells = insurance_table.rows[0].cells
        header_cells[0].paragraphs[0].add_run('Brief description of policy including policy number').bold = True
        header_cells[1].paragraphs[0].add_run(
            'Insurance company or underwriter (including address for service of notices)').bold = True
        header_cells[2].paragraphs[0].add_run('Date of Policy').bold = True

        # Add empty row for completion
        row = insurance_table.add_row()

    @classmethod
    def _add_schedule_3(cls, doc, context):
        """Add Schedule 3"""
        doc.add_page_break()

        # Schedule 3 header
        schedule_header = doc.add_paragraph()
        schedule_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        schedule_run = schedule_header.add_run('SCHEDULE 3')
        schedule_run.bold = True
        schedule_run.font.size = Pt(14)
        schedule_run.underline = True
        schedule_header.paragraph_format.space_after = Pt(12)

        # Part I
        part1_header = doc.add_paragraph()
        part1_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part1_run = part1_header.add_run('Part I\nForm of notice relating to assigned Occupational Leases')
        part1_run.bold = True
        part1_run.font.size = Pt(12)
        part1_header.paragraph_format.space_after = Pt(12)

        # Add notice form content
        cls._add_occupational_lease_notice_form(doc, context)

        # Part II
        doc.add_page_break()
        part2_header = doc.add_paragraph()
        part2_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part2_run = part2_header.add_run('Part II\nAcknowledgement of Notice of Assignment of Occupational Leases')
        part2_run.bold = True
        part2_run.font.size = Pt(12)
        part2_header.paragraph_format.space_after = Pt(12)

        # Add acknowledgement form
        cls._add_occupational_lease_acknowledgement_form(doc, context)

    @classmethod
    def _add_schedule_4(cls, doc, context):
        """Add Schedule 4"""
        doc.add_page_break()

        # Schedule 4 header
        schedule_header = doc.add_paragraph()
        schedule_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        schedule_run = schedule_header.add_run('SCHEDULE 4')
        schedule_run.bold = True
        schedule_run.font.size = Pt(14)
        schedule_run.underline = True
        schedule_header.paragraph_format.space_after = Pt(12)

        # Part I
        part1_header = doc.add_paragraph()
        part1_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part1_run = part1_header.add_run('Part I\nForm of notice relating to Insurances')
        part1_run.bold = True
        part1_run.font.size = Pt(12)
        part1_header.paragraph_format.space_after = Pt(12)

        # Add insurance notice form
        cls._add_insurance_notice_form(doc, context)

        # Part II
        doc.add_page_break()
        part2_header = doc.add_paragraph()
        part2_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part2_run = part2_header.add_run('Part II\nAcknowledgement of Notice of Assignment of Insurances')
        part2_run.bold = True
        part2_run.font.size = Pt(12)
        part2_header.paragraph_format.space_after = Pt(12)

        # Add insurance acknowledgement form
        cls._add_insurance_acknowledgement_form(doc, context)

    @classmethod
    def _add_execution_page(cls, doc, context):
        """Add execution page"""
        doc.add_page_break()

        # IN WITNESS header
        witness_para = doc.add_paragraph()
        witness_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        witness_run = witness_para.add_run('IN WITNESS ')
        witness_run.bold = True
        witness_run.font.size = Pt(12)
        witness_para.add_run(
            'whereof this Deed has been executed and delivered as a deed on the date first written above.')
        witness_para.paragraph_format.space_after = Pt(24)

        # Execution page title
        exec_title = doc.add_paragraph()
        exec_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        exec_title_run = exec_title.add_run('EXECUTION PAGE')
        exec_title_run.bold = True
        exec_title_run.font.size = Pt(14)
        exec_title.paragraph_format.space_after = Pt(24)

        # Create execution table
        exec_table = doc.add_table(rows=2, cols=1)
        exec_table.style = 'Table Grid'
        exec_table.columns[0].width = Inches(7.0)

        # Chargor execution section
        chargor_cell = exec_table.cell(0, 0)
        cls._add_execution_section(chargor_cell, 'THE CHARGOR', context.get('chargor_name', ''))

        # Lender execution section
        lender_cell = exec_table.cell(1, 0)
        cls._add_execution_section(lender_cell, 'THE LENDER', context.get('lender_name', ''))

    # HELPER METHODS FOR FORMATTING

    @classmethod
    def _add_main_clause_header(cls, doc, clause_num, clause_title):
        """Add main clause header with consistent formatting"""
        clause_para = doc.add_paragraph()
        clause_para.add_run(f'{clause_num}. ').bold = True
        title_run = clause_para.add_run(clause_title)
        title_run.bold = True
        title_run.underline = True
        title_run.font.size = Pt(12)
        title_run.font.name = 'Times New Roman'
        clause_para.paragraph_format.space_after = Pt(12)
        clause_para.paragraph_format.space_before = Pt(12)

    @classmethod
    def _add_subsection_header(cls, doc, subsection_num, subsection_title):
        """Add subsection header"""
        subsection_para = doc.add_paragraph()
        subsection_para.paragraph_format.left_indent = Inches(0.5)
        subsection_para.add_run(f'{subsection_num} ').bold = True
        subsection_para.add_run(subsection_title).bold = True
        subsection_para.paragraph_format.space_after = Pt(8)

    @classmethod
    def _add_subsection_content(cls, doc, content):
        """Add subsection content"""
        content_para = doc.add_paragraph()
        content_para.paragraph_format.left_indent = Inches(0.5)
        content_para.paragraph_format.space_after = Pt(8)
        content_para.paragraph_format.line_spacing = 1.15
        content_para.add_run(content)

    @classmethod
    def _add_numbered_subsection(cls, doc, num, content):
        """Add numbered subsection"""
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.75)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.15

        num_run = para.add_run(f'{num} ')
        num_run.bold = True
        para.add_run(content)

    @classmethod
    def _add_numbered_list_item(cls, doc, num, content):
        """Add numbered list item"""
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(1.0)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.1

        num_run = para.add_run(f'{num}. ')
        num_run.bold = True
        para.add_run(content)

    @classmethod
    def _add_numbered_sublist_item(cls, doc, letter, content):
        """Add numbered sublist item with letter"""
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(1.5)
        para.paragraph_format.first_line_indent = Inches(-0.25)

        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.1

        letter_run = para.add_run(f'({letter}) ')
        letter_run.bold = True
        para.add_run(content)

    @classmethod
    def _add_numbered_event_of_default(cls, doc, num, event_name, event_desc):
        """Add numbered event of default with special formatting"""
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(1.0)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.1

        num_run = para.add_run(f'{num}. ')
        num_run.bold = True
        event_name_run = para.add_run(f'{event_name}: ')
        event_name_run.bold = True
        para.add_run(event_desc)

    @classmethod
    def _add_definition(cls, doc, term, definition):
        """Add definition with proper formatting"""
        def_para = doc.add_paragraph()
        def_para.paragraph_format.left_indent = Inches(0.75)
        def_para.paragraph_format.first_line_indent = Inches(-0.25)
        def_para.paragraph_format.space_after = Pt(6)
        def_para.paragraph_format.line_spacing = 1.1

        term_run = def_para.add_run(term)
        term_run.bold = True
        def_para.add_run(f' {definition}')

    @classmethod
    def _add_construction_subsections(cls, doc):
        """Add construction subsections for clause 1.4"""
        # 1.4.1
        cls._add_numbered_subsection(doc, '1.4.1',
                                     'References to this Deed or to any other agreement or document shall be construed as a reference to this Deed or, as the case may be, such other agreement or document as the same may have been, or may from time to time be, amended, restated, varied, novated, assigned, substituted, supplemented or otherwise modified from time to time (and so that any reference to this Deed shall include, unless the context otherwise requires each Prescribed Form Charge and any other agreement or document expressed to be supplemental hereto or expressed to be collateral herewith or which is otherwise entered into pursuant to or in accordance with the provisions hereof).')

        # 1.4.2
        cls._add_numbered_subsection(doc, '1.4.2',
                                     'In this Deed (including the Recitals), all terms and expressions shall, unless otherwise defined in this Deed or the context requires otherwise, have the meaning attributed to such terms in the Loan Agreement (whether defined expressly in the Loan Agreement or by reference to another document).')

        # 1.4.3
        cls._add_numbered_subsection(doc, '1.4.3',
                                     'Unless a contrary indication appears in this Deed:')

        # Sub-items for 1.4.3
        construction_items = [
            'words and phrases the definition of which is contained in or referred to in section 2 of the Companies Act 2014 are to be construed as having the meaning attributed to them in that section;',
            'references to any enactments or other legislation shall be deemed to include references to such enactment or other legislation as re-enacted, amended, substituted or extended from time to time;',
            'references to clauses and Schedules are to be construed as references to the clauses of and the Schedules to this Deed and any reference to this Deed includes each of its Schedules;',
            'words importing the plural shall include the singular and vice versa and words denoting any gender include all genders.'
        ]

        for i, item in enumerate(construction_items, 1):
            cls._add_numbered_sublist_item(doc, str(i), item)

        # Continue with remaining construction subsections...
        # 1.4.4, 1.4.5, 1.4.6, etc.

    @classmethod
    def _add_address_table(cls, doc, context):
        """Add address table for notices section"""
        # Create address table
        addr_table = doc.add_table(rows=2, cols=2)
        addr_table.style = 'Table Grid'
        addr_table.columns[0].width = Inches(1.5)
        addr_table.columns[1].width = Inches(5.5)

        # Chargor row
        chargor_label_cell = addr_table.cell(0, 0)
        chargor_label_cell.paragraphs[0].add_run('The Chargor:').bold = True
        chargor_label_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        chargor_details_cell = addr_table.cell(0, 1)
        chargor_details = chargor_details_cell.paragraphs[0]
        chargor_details.add_run('Address: ').bold = True
        chargor_details.add_run('_' * 50)
        chargor_details.add_run('\n\nAttention: ')
        chargor_details.add_run('_' * 40)
        chargor_details.add_run('\n\nEmail Address: ')
        chargor_details.add_run('_' * 40)

        # Lender row
        lender_label_cell = addr_table.cell(1, 0)
        lender_label_cell.paragraphs[0].add_run('The Lender:').bold = True
        lender_label_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        lender_details_cell = addr_table.cell(1, 1)
        lender_details = lender_details_cell.paragraphs[0]
        lender_details.add_run('Address: ').bold = True
        lender_details.add_run('_' * 50)
        lender_details.add_run('\n\nAttention: ')
        lender_details.add_run('_' * 40)
        lender_details.add_run('\n\nEmail Address: ')
        lender_details.add_run('_' * 40)

    @classmethod
    def _add_form_52_content(cls, doc, context):
        """Add Form 52 content with professional table layout"""
        form_title = doc.add_paragraph()
        form_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        form_title_run = form_title.add_run(
            'Specific charge for present and future advances arising on the creation of a commercial mortgage or debenture, (rules 52, 105)')
        form_title_run.italic = True
        form_title_run.font.size = Pt(10)
        form_title_run.font.name = 'Times New Roman'
        form_title.paragraph_format.space_after = Pt(12)

        land_registry_title = doc.add_paragraph()
        land_registry_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        land_registry_run = land_registry_title.add_run('LAND REGISTRY\n\nSPECIFIC CHARGE')
        land_registry_run.bold = True
        land_registry_run.font.size = Pt(14)
        land_registry_run.font.name = 'Times New Roman'
        land_registry_title.paragraph_format.space_after = Pt(16)

        # Create form table
        form_table = doc.add_table(rows=1, cols=2)
        form_table.style = 'Table Grid'
        form_table.columns[0].width = Inches(2.5)
        form_table.columns[1].width = Inches(4.5)

        # Remove default row
        form_table._element.remove(form_table.rows[0]._element)

        # Add form fields
        form_fields = [
            ('Date:', ''),
            ('Secured Party:', ''),
            ('Mortgagor:', ''),
            ('Mortgaged Property subject to specific charge:',
             f"The property comprised in Folio {context.get('folio_number', '_' * 20)} County {context.get('county', '_' * 20)}\n\nALL THAT the property known as {context.get('property_address', '_' * 40)}\n\n(use a continuation sheet if necessary)"),
            ('Mortgage Conditions:',
             'This Mortgage incorporates the Mortgage Conditions in Mortgage/Debenture of even date, between the parties herein, as if they were set out in this Mortgage in full. The term "Secured Liabilities" has the meaning given in the Mortgage Conditions.'),
            ('SPECIFIC CHARGE:',
             'As security for the payment and discharge of the Secured Liabilities, the Mortgagor as beneficial owner (and also in the case of registered land as registered owner or as the person entitled to be registered as registered owner) hereby charges in favour of the Secured Party the Mortgaged Property with the payment of the Secured Liabilities, and assents to the registration of this charge as a burden on the Mortgaged Property.\n\nThe Mortgagor acknowledges that the charge hereby created forms one transaction with the security created in the aforesaid Mortgage/Debenture of even date for payment of the Secured Liabilities.'),
        ]

        for label, content in form_fields:
            row = form_table.add_row()
            label_cell = row.cells[0]
            content_cell = row.cells[1]

            # Format label cell
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            label_run.font.name = 'Times New Roman'
            label_run.font.size = Pt(11)
            label_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            # Format content cell
            content_para = content_cell.paragraphs[0]
            if content:
                content_run = content_para.add_run(content)
                content_run.font.name = 'Times New Roman'
                content_run.font.size = Pt(10)
            content_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Add signatures section
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        sig_para = doc.add_paragraph()
        sig_run = sig_para.add_run('Signatures:')
        sig_run.bold = True
        sig_run.font.name = 'Times New Roman'
        sig_run.font.size = Pt(11)
        sig_para.paragraph_format.space_after = Pt(8)

        # Create signatures table
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.style = 'Table Grid'
        sig_table.columns[0].width = Inches(3.5)
        sig_table.columns[1].width = Inches(3.5)

        # Left signature section
        left_cell = sig_table.cell(0, 0)
        cls._add_signature_section(left_cell)

        # Right signature section
        right_cell = sig_table.cell(0, 1)
        cls._add_signature_section(right_cell)

        # Add note
        note = doc.add_paragraph()
        note_run = note.add_run(
            'Note - For execution and the attestation of the execution of a charge - see Rules 54 and 55.')
        note_run.bold = True
        note_run.font.size = Pt(9)
        note_run.font.name = 'Times New Roman'

    @classmethod
    def _add_signature_section(cls, cell):
        """Add signature section to a cell"""
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0

        para.add_run('Present when the common seal of').font.size = Pt(9)
        para.add_run('\n\n')
        para.add_run('_' * 35).font.size = Pt(9)
        para.add_run('\n\n')
        para.add_run('Was affixed to this deed and this deed was delivered').font.size = Pt(9)
        para.add_run('\n\n')
        para.add_run('_' * 35).font.size = Pt(9)
        para.add_run('\n')
        para.add_run('Director').font.size = Pt(9)
        para.add_run('\n\n')
        para.add_run('_' * 35).font.size = Pt(9)
        para.add_run('\n')
        para.add_run('Director/Secretary').font.size = Pt(9)

    @classmethod
    def _add_occupational_lease_notice_form(cls, doc, context):
        """Add occupational lease notice form"""
        # From/To/Date section
        from_para = doc.add_paragraph()
        from_para.add_run('From: ').bold = True
        from_para.add_run('[CHARGOR] (the "Chargor")')
        from_para.paragraph_format.space_after = Pt(8)

        to_para = doc.add_paragraph()
        to_para.add_run('To: ').bold = True
        to_para.add_run('[COUNTERPARTY] (the "Contractual Party")')
        to_para.paragraph_format.space_after = Pt(8)

        date_para = doc.add_paragraph()
        date_para.add_run('Date: ').bold = True
        date_para.add_run('20[ ]')
        date_para.paragraph_format.space_after = Pt(12)

        # Re: section
        re_para = doc.add_paragraph()
        re_para.add_run('Re: ').bold = True
        re_para.add_run('[INSERT CONTRACT/LEASE DESCRIPTION] ')
        re_para.add_run('between [(1) the Chargor and (2) the Contractual Party] ').bold = True
        re_para.add_run('[INSERT CORRECT PARTY DETAILS AS APPROPRIATE] ')
        re_para.add_run('OR ')
        re_para.add_run('[(1) the Contractual Party and (2) the Chargor] ').bold = True
        re_para.add_run('[INSERT CORRECT PARTY DETAILS AS APPROPRIATE] (the ')
        re_para.add_run('"Agreement"').bold = True
        re_para.add_run(').')
        re_para.paragraph_format.space_after = Pt(12)

        # Dear Sirs
        dear_para = doc.add_paragraph()
        dear_para.add_run('Dear Sirs')
        dear_para.paragraph_format.space_after = Pt(12)

        # Main notice content
        notice_content = '''We hereby give you notice that we have assigned by way of security to [●] (the "Lender", which term shall include its successors and assigns) pursuant to a Mortgage and Charge dated [●] 20[ ] entered into by us in favour of the Lender (the "Mortgage and Charge") all our right, title and interest in (but not our obligations) to and under the Agreement.
    
    We confirm that:
    
    (a) for the avoidance of doubt, we will remain liable under the Agreement to perform all the obligations assumed by us under the Agreement; and
    
    (b) at no time will the Lender, any of its agents, any Receiver nor any other person be under any obligation or liability to you under or in respect of the Agreement.
    
    We remain entitled to exercise all our rights, powers and discretions under the Agreement and you should continue to give notices under the Agreement to us, unless and until you receive notice (the "Default Notice") from the Lender to the contrary stating that the security constituted by the Mortgage and Charge has become enforceable.
    
    Immediately following receipt of a Default Notice:
    
    (a) all payments from you under or arising from the Agreement should be made to such account(s) as may from time to time be notified to you in writing by the Lender;
    
    (b) all remedies provided for in the Agreement or available at law or in equity are exercisable by the Lender;
    
    (c) all rights to compel performance of the Agreement are exercisable by the Lender; and
    
    (d) all rights, interests and benefits whatsoever accruing to or for the benefit of ourselves arising from the Agreement belong to the Lender.
    
    This notice and the terms set out in this notice shall be irrevocable save as otherwise advised in writing by the Lender. This notice shall be governed by and construed with the laws of Ireland.
    
    Please acknowledge receipt of this notice by signing the acknowledgement on the enclosed copy letter and returning same to the Lender at [INSERT ADDRESS] marked for the attention of [INSERT DETAILS OF LENDER(S)].
    
    Yours faithfully
    
    _________________________________
    
    for and on behalf of
    
    [CHARGOR]'''

        content_para = doc.add_paragraph()
        content_para.add_run(notice_content)
        content_para.paragraph_format.line_spacing = 1.15
        content_para.paragraph_format.space_after = Pt(12)

    @classmethod
    def _add_occupational_lease_acknowledgement_form(cls, doc, context):
        """Add occupational lease acknowledgement form"""
        acknowledgement_content = '''To: [INSERT LENDER NAME AND ADDRESS]
    
    Attention: [●]
    
    Date: [•]
    
    Re: [INSERT CONTRACT/LEASE DESCRIPTION] between [(1) the Chargor and (2) the Contractual Party] [INSERT CORRECT PARTY DETAILS AS APPROPRIATE] OR [(1) the Contractual Party and (2) the Chargor] [INSERT CORRECT PARTY DETAILS AS APPROPRIATE] (the "Agreement").
    
    Dear Sirs
    
    We acknowledge receipt of a notice in the terms attached (the "Notice") and confirm our consent to the assignment and charge referred to in the Notice. We further confirm that we have not received notice of any previous assignments or charges of or over any of the rights, interests and benefits in and to the Agreement as referred to in the Notice. Terms defined in the Notice shall have the same meanings herein.
    
    With effect from the receipt by us of the Default Notice (as defined in the Notice), we shall pay all monies due by us under or arising from the Agreement in the manner specified in the Notice.
    
    We further confirm that no amendment, waiver or release of any such rights, interests and benefits shall be effective without your prior written consent. Furthermore we confirm that no breach or default on the part of the Chargor of any of the terms of the Agreement shall be deemed to have occurred unless we have given notice of such breach to you specifying how to make good such breach.
    
    We also confirm that we shall not exercise any right of combination, consolidation or set--off which we may have in respect of any amount due under the Agreement.
    
    _________________________________
    
    for and on behalf of
    
    [INSERT DETAILS OF COUNTERPARTY]'''

        content_para = doc.add_paragraph()
        content_para.add_run(acknowledgement_content)
        content_para.paragraph_format.line_spacing = 1.15

    @classmethod
    def _add_insurance_notice_form(cls, doc, context):
        """Add insurance notice form"""
        insurance_notice_content = '''From: [CHARGOR] Limited
    
    To: [INSURANCE COMPANY]
    
    Date: 20[ ]
    
    Re: [INSERT DETAILS OF INSURANCE POLICY] (the "Policy")
    
    Dear Sirs
    
    We hereby give you notice that we have assigned by way of security to [LENDER] (the "Lender", which term shall include its successors and assigns) pursuant to a Mortgage and Charge dated [●] 20[ ] entered into by us in favour of the Lender (the "Mortgage and Charge") all our right, title and interest in, to and under the Policy including all monies payable under the Policy, proceeds of all claims, awards and judgments and all other insurances entered into supplemental to or in replacement of such Policy.
    
    We will remain liable to perform all our obligations under the Policy and the Lender is under no obligation of any kind whatsoever under the Policy nor under any liability whatsoever in the event of any failure by us to perform our obligations under the Policy.
    
    We irrevocably instruct and authorise you, after receipt of this notice, to make all payments under or arising under the Policy to such accounts as may from time to time be notified to you by the Lender.
    
    Please note that all rights, interests and benefits whatsoever accruing to or for the benefit of ourselves arising from the Policy belong to the Lender.
    
    We hereby instruct you to note the interest of the Lender on the Policy as [JOINT INSURED/INTERESTED PARTY/SOLE LOSS PAYEE] and authorise you to disclose to the Lender, without further approval from us, such information regarding the Policy as the Lender may from time to time request and to send it copies of all notices issued by you under the Policy.
    
    This notice and the terms set out herein shall be irrevocable save as otherwise advised in writing by the Lender and shall be governed by and construed with the laws of Ireland.
    
    Please acknowledge receipt of this notice by signing the acknowledgement on the enclosed copy letter and returning same to the Lender at [INSERT ADDRESS] marked for the attention of [INSERT DETAILS OF LENDER(S)].
    
    Yours faithfully
    
    _________________________________
    
    for and on behalf of
    
    [CHARGOR]'''

        content_para = doc.add_paragraph()
        content_para.add_run(insurance_notice_content)
        content_para.paragraph_format.line_spacing = 1.15

    @classmethod
    def _add_insurance_acknowledgement_form(cls, doc, context):
        """Add insurance acknowledgement form"""
        insurance_ack_content = '''To: [LENDER]
    
    Date: 20[ ]
    
    Re: [INSERT DETAILS OF INSURANCE POLICY] (the "Policy")
    
    Dear Sirs
    
    We confirm receipt from [CHARGOR] (the "Chargor") of a notice dated [●] 20[ ] (the "Notice") of an assignment upon the terms of a Mortgage and Charge dated [●] 20[ ] (the "Mortgage and Charge") between (1) the Chargor and (2) you (the "Lender") in respect of all the Chargor's rights, title, interest and benefit in, to and under the Policy (as specified in the Notice).
    
    We confirm that we have not received notice of any assignment or charge of or over any of the right, interests and benefits specified in the Notice and will make all payments to any account as may from time to time be specified by you.
    
    We acknowledge that the Chargor will remain liable to perform all of its obligations under the Policy and the Lender is under no obligation of any kind whatsoever under the Policy nor under any liability whatsoever in the event of any failure by us to perform our obligations under the Policy.
    
    This letter is governed by and will be construed in accordance with the laws of Ireland.
    
    Yours faithfully
    
    _________________________________
    
    [Insert name of insurance company]'''

        content_para = doc.add_paragraph()
        content_para.add_run(insurance_ack_content)
        content_para.paragraph_format.line_spacing = 1.15

    @classmethod
    def _add_execution_section(cls, cell, party_title, party_name):
        """Add execution section for a party"""
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)

        # Party title
        title_run = para.add_run(party_title)
        title_run.bold = True
        title_run.underline = True
        title_run.font.size = Pt(12)
        para.add_run('\n\n')

        # Signed by
        para.add_run('SIGNED by ').bold = True
        para.add_run(f'{party_name or "_" * 30}')
        para.add_run('\n\n\n')
        para.add_run('_' * 50)
        para.add_run('\n\n')

        # Witness details
        para.add_run('Witness Signature: ').bold = True
        para.add_run('\n\n')
        para.add_run('_' * 50)
        para.add_run('\n\n')

        para.add_run('Witness Name: ').bold = True
        para.add_run('\n\n')
        para.add_run('_' * 50)
        para.add_run('\n\n')

        para.add_run('Witness Address: ').bold = True
        para.add_run('\n\n')
        para.add_run('_' * 50)
        para.add_run('\n\n')
        para.add_run('_' * 50)
        para.add_run('\n\n')
        para.add_run('_' * 50)
        para.add_run('\n\n')

        para.add_run('Witness Occupation: ').bold = True
        para.add_run('\n\n')
        para.add_run('_' * 50)
