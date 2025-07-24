# document_requirements/services.py - Complete Document Generation Service

from django.template.loader import render_to_string
from django.utils import timezone
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt
from xhtml2pdf import pisa
from io import BytesIO
import logging
import os

logger = logging.getLogger(__name__)


class DocumentTemplateService:
    """Service for generating PDF/Word documents on-the-fly from HTML templates or serving static files"""

    @classmethod
    def can_generate_template(cls, document_type):
        supported_templates = [
            "Beneficiaries Irrevocable Instruction to Law Firm",
            "Solicitor Letter of Undertaking (Not to distribute estate)",
            "Renunciation of Probate",
            "Certificate of Title",
            "Land Registry Form 51",
        ]

        if not document_type.has_template:
            return False
        return any(template_name in document_type.name for template_name in supported_templates)

    @classmethod
    def generate_document_response(cls, requirement):
        """Generate document (PDF/Word) and return BytesIO response for download"""
        try:
            if not requirement.document_type.has_template:
                logger.warning(f"Document type {requirement.document_type.name} does not have template enabled")
                return None

            # Handle different document types
            if "Beneficiaries Irrevocable Instruction to Law Firm" in requirement.document_type.name:
                return cls._generate_beneficiaries_authorisation_pdf(requirement)
            elif "Solicitor Letter of Undertaking (Not to distribute estate)" in requirement.document_type.name:
                return cls._generate_solicitors_letter_of_undertaking(requirement)
            elif "Renunciation of Probate" in requirement.document_type.name:
                return cls._generate_renunciation_word_document(requirement)
            elif "Certificate of Title" in requirement.document_type.name:
                return cls._serve_static_certificate_of_title(requirement)
            elif "Land Registry Form 51" in requirement.document_type.name:
                return cls._generate_land_registry_form_51(requirement)
            else:
                logger.warning(f"Template generation not implemented for: {requirement.document_type.name}")
                return None

        except Exception as e:
            logger.error(f"Error generating template document: {str(e)}")
            return None

    @classmethod
    def _generate_beneficiaries_authorisation_pdf(cls, requirement):
        """Generate the Beneficiaries Authorisation PDF from HTML template"""
        try:
            context = cls._get_template_context(requirement)
            html_string = render_to_string('document_templates/beneficiaries_authorisation.html', context)

            result = BytesIO()
            pdf = pisa.CreatePDF(BytesIO(html_string.encode("UTF-8")), dest=result)

            if pdf.err:
                logger.error("Error generating Beneficiaries Authorisation PDF")
                return None

            return result

        except Exception as e:
            logger.error(f"Error generating Beneficiaries Authorisation PDF: {str(e)}")
            return None

    @classmethod
    def _generate_solicitors_letter_of_undertaking(cls, requirement):
        """Generate the Solicitor Letter of Undertaking PDF from HTML template"""
        try:
            context = cls._get_template_context(requirement)
            html_string = render_to_string('document_templates/solicitor_letter_of_undertaking.html', context)

            result = BytesIO()
            pdf = pisa.CreatePDF(BytesIO(html_string.encode("UTF-8")), dest=result)

            if pdf.err:
                logger.error("Error generating Solicitor Letter of Undertaking PDF")
                return None

            return result

        except Exception as e:
            logger.error(f"Error generating Solicitor Letter of Undertaking PDF: {str(e)}")
            return None

    @classmethod
    def _generate_renunciation_word_document(cls, requirement):
        """Generate the Renunciation of Probate Word document"""
        try:
            # Import docx libraries
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError as e:
                logger.error(f"python-docx library not installed: {str(e)}")
                logger.error("Please install with: pip install python-docx")
                return None

            # Get application data
            application = requirement.application
            deceased = application.deceased if hasattr(application, 'deceased') else None
            logger.info(f"Processing renunciation for application {application.id}")

            # Create Word document
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(1.2)
                section.right_margin = Inches(1.2)

            # Title
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run('RENUNCIATION OF PROBATE')
            title_run.bold = True
            title_run.underline = True
            title_run.font.size = Pt(16)
            title_run.font.name = 'Times New Roman'
            title.paragraph_format.space_after = Pt(24)

            # Estate title
            estate_title = doc.add_paragraph()
            estate_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            estate_title.add_run('In the estate of ')

            if deceased and hasattr(deceased, 'first_name') and hasattr(deceased, 'last_name'):
                name_run = estate_title.add_run(f'{deceased.first_name} {deceased.last_name}')
                name_run.font.name = 'Times New Roman'
                name_run.bold = True
            else:
                name_run = estate_title.add_run('................................................')
                name_run.underline = True

            estate_title.add_run(' (named the deceased)')
            estate_title.paragraph_format.space_after = Pt(18)

            # Main paragraph 1
            para1 = doc.add_paragraph()
            para1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para1_format = para1.paragraph_format
            para1_format.line_spacing = 1.15
            para1_format.space_after = Pt(12)
            para1_format.first_line_indent = Inches(0.25)

            # Build first paragraph with Irish legal formatting
            para1.add_run('Whereas ')

            if deceased and hasattr(deceased, 'first_name') and hasattr(deceased, 'last_name'):
                name_run = para1.add_run(f'{deceased.first_name} {deceased.last_name}')
                name_run.bold = True
            else:
                name_run = para1.add_run('................................................')
                name_run.underline = True

            para1.add_run(', late of ')

            if deceased and hasattr(deceased, 'address') and deceased.address:
                addr_run = para1.add_run(str(deceased.address))
                addr_run.bold = True
            else:
                addr_run = para1.add_run(
                    '................................................................................................................................................................................................')
                addr_run.underline = True

            para1.add_run(' deceased, died on the ')

            if deceased and hasattr(deceased, 'date_of_death') and deceased.date_of_death:
                day_run = para1.add_run(str(deceased.date_of_death.day))
                day_run.bold = True
                para1.add_run(' day of ')
                my_run = para1.add_run(deceased.date_of_death.strftime('%B %Y'))
                my_run.bold = True
            else:
                day_run = para1.add_run('........')
                day_run.underline = True
                para1.add_run(' day of ')
                month_run = para1.add_run('........................ 20......')
                month_run.underline = True

            para1.add_run(
                ', at [where the application is made in District Probate Registry, add having at the time of his / her death a fixed abode at ')

            fixed_addr_run = para1.add_run(
                '................................................................................................................................................................................................')
            fixed_addr_run.underline = True

            para1.add_run(' within the district of ')

            district_run = para1.add_run('......................................')
            district_run.underline = True

            para1.add_run(
                '] and whereas he/she made and duly executed his/her last will [or will and codicils] bearing date the ')

            will_day_run = para1.add_run('........')
            will_day_run.underline = True

            para1.add_run(' day of ')

            will_month_run = para1.add_run('........................ 20......')
            will_month_run.underline = True

            para1.add_run(', and thereof naming ')

            executor_run = para1.add_run('................................................')
            executor_run.underline = True

            para1.add_run(' sole executor.')

            # Main paragraph 2
            para2 = doc.add_paragraph()
            para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para2_format = para2.paragraph_format
            para2_format.line_spacing = 1.15
            para2_format.space_after = Pt(24)
            para2_format.first_line_indent = Inches(0.25)

            para2.add_run('Now I, the said ')
            executor_name_run = para2.add_run('................................................')
            executor_name_run.underline = True
            para2.add_run(
                ', aged 18 years and upwards, do declare that I have not intermeddled with the estate of the said deceased, and will not hereafter intermeddle therein with the intent to defraud creditors, and I do expressly renounce my right to probate of the said will and of the estate of the said deceased.')

            # Signature section
            dated_para = doc.add_paragraph()
            dated_para.paragraph_format.space_before = Pt(12)
            dated_para.paragraph_format.space_after = Pt(8)
            dated_run = dated_para.add_run('Dated: ')
            dated_run.font.name = 'Times New Roman'
            dated_run.font.size = Pt(12)
            dated_run.bold = True

            date_line = doc.add_paragraph()
            date_line.add_run(
                '................................................................................................................................................................................................')
            date_line.paragraph_format.space_after = Pt(16)

            signed_para = doc.add_paragraph()
            signed_para.paragraph_format.space_after = Pt(8)
            signed_run = signed_para.add_run('Signed: ')
            signed_run.font.name = 'Times New Roman'
            signed_run.font.size = Pt(12)
            signed_run.bold = True

            sig_line = doc.add_paragraph()
            sig_line.add_run(
                '................................................................................................................................................................................................')
            sig_line.paragraph_format.space_after = Pt(20)

            # Witness section
            witness_title = doc.add_paragraph()
            witness_title.paragraph_format.space_before = Pt(8)
            witness_run = witness_title.add_run('WITNESS')
            witness_run.font.name = 'Times New Roman'
            witness_run.font.size = Pt(12)
            witness_run.bold = True
            witness_run.underline = True
            witness_title.paragraph_format.space_after = Pt(12)

            # Witness fields
            witness_name_label = doc.add_paragraph()
            witness_name_label.add_run('Name: ').bold = True
            witness_name_label.paragraph_format.space_after = Pt(4)

            witness_name_line = doc.add_paragraph()
            witness_name_line.add_run(
                '................................................................................................................................................................................................')
            witness_name_line.paragraph_format.space_after = Pt(10)

            witness_addr_label = doc.add_paragraph()
            witness_addr_label.add_run('Address: ').bold = True
            witness_addr_label.paragraph_format.space_after = Pt(4)

            for i in range(3):
                addr_line = doc.add_paragraph()
                addr_line.add_run(
                    '................................................................................................................................................................................................')
                addr_line.paragraph_format.space_after = Pt(6)

            witness_sig_label = doc.add_paragraph()
            witness_sig_label.add_run('Witness Signature: ').bold = True
            witness_sig_label.paragraph_format.space_before = Pt(8)
            witness_sig_label.paragraph_format.space_after = Pt(4)

            witness_sig_line = doc.add_paragraph()
            witness_sig_line.add_run(
                '................................................................................................................................................................................................')

            # Save to BytesIO
            result = BytesIO()
            doc.save(result)
            result.seek(0)
            logger.info("Renunciation Word document created successfully")
            return result

        except Exception as e:
            logger.error(f"Error generating Renunciation Word document: {str(e)}")
            import traceback
            logger.error(f"Full traceback: {traceback.format_exc()}")
            return None

    @classmethod
    def _generate_land_registry_form_51(cls, requirement):
        """Generate the Land Registry Form 51 Word document"""
        try:
            # Import docx libraries
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.table import WD_ALIGN_VERTICAL
            except ImportError as e:
                logger.error(f"python-docx library not installed: {str(e)}")
                logger.error("Please install with: pip install python-docx")
                return None

            # Get application data
            application = requirement.application
            context = cls._get_land_registry_context(requirement)
            logger.info(f"Processing Land Registry Form 51 for application {application.id}")

            # Create Word document
            doc = Document()

            # Set tight margins to fit more content
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)

            # Form header - compact
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_run = header.add_run('FORM 51')
            header_run.bold = True
            header_run.font.size = Pt(14)
            header_run.font.name = 'Times New Roman'
            header.paragraph_format.space_after = Pt(3)

            # Subtitle - compact
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.add_run('Charge for present and future advances (Rules 52 and 105)')
            subtitle_run.italic = True
            subtitle_run.font.size = Pt(9)
            subtitle_run.font.name = 'Times New Roman'
            subtitle.paragraph_format.space_after = Pt(6)

            # Land Registry title - compact
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run('LAND REGISTRY')
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_run.font.name = 'Times New Roman'
            title.paragraph_format.space_after = Pt(3)

            # Mortgage subtitle - compact
            mortgage_title = doc.add_paragraph()
            mortgage_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mortgage_run = mortgage_title.add_run('MORTGAGE')
            mortgage_run.bold = True
            mortgage_run.font.size = Pt(12)
            mortgage_run.font.name = 'Times New Roman'
            mortgage_title.paragraph_format.space_after = Pt(8)

            # Create main table - more compact
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'

            # Set tighter column widths
            table.columns[0].width = Inches(1.8)
            table.columns[1].width = Inches(5.2)

            # Remove the default row
            table._element.remove(table.rows[0]._element)

            # Date row
            cls._add_form_row(table, 'Date:', '', fill_type='blank')

            # Secured Party row
            cls._add_form_row(table, 'Secured Party:', '', fill_type='blank')

            # Mortgagor row
            cls._add_form_row(table, 'Mortgagor:', '', fill_type='blank')

            # Mortgaged Property row (more compact)
            property_content = f"""The property comprised in Folio {context.get('folio_number', '____________________')} County {context.get('county', '____________________')}

ALL THAT the property known as {context.get('property_address', '_' * 60)}

(*use a continuation sheet if necessary*)"""
            cls._add_form_row(table, 'Mortgaged Property:', property_content, is_large=True, compact=True)

            # General Mortgage Conditions row (more compact)
            conditions_content = """This Mortgage incorporates the Loan Mortgage Conditions as if they were set out in this Mortgage in full and the Mortgagor acknowledges that the Mortgagor has been given a copy of the General Mortgage Conditions and has read them and agrees to be bound by them. The term 'Secured Liabilities' has the meaning given in the 'General Conditions'."""
            cls._add_form_row(table, 'General Mortgage Conditions:', conditions_content, is_large=True, compact=True)

            # Mortgage section (more compact)
            mortgage_content = """As security for the payment and discharge of the Secured Liabilities, the Mortgagor as beneficial owner (and also in the case of registered land as registered owner or as the person entitled to be registered as registered owner) hereby charges in favour of the Secured Party the Mortgaged Property with the payment of the Secured Liabilities, and assents to the registration of this charge as a burden on the Mortgaged Property."""
            cls._add_form_row(table, 'Mortgage:', mortgage_content, is_large=True, compact=True)

            # Add minimal space before signatures
            doc.add_paragraph().paragraph_format.space_after = Pt(8)

            # Create signatures section with connected header
            sig_header = doc.add_paragraph()
            sig_header_run = sig_header.add_run('Signatures:')
            sig_header_run.bold = True
            sig_header_run.font.name = 'Times New Roman'
            sig_header_run.font.size = Pt(11)
            sig_header.paragraph_format.space_after = Pt(6)

            # Create compact signatures table
            sig_table = doc.add_table(rows=1, cols=2)
            sig_table.style = 'Table Grid'

            # Set signature table column widths
            sig_table.columns[0].width = Inches(3.5)
            sig_table.columns[1].width = Inches(3.5)

            # Left column - First signatory (proper signature spacing)
            left_cell = sig_table.cell(0, 0)
            left_para = left_cell.paragraphs[0]
            left_para.paragraph_format.space_after = Pt(0)
            left_para.paragraph_format.line_spacing = 1.0

            # Main signature section with substantial space
            left_para.add_run('Signed and Delivered as a deed:').bold = True
            left_para.add_run('\n\n')  # 5 line breaks for signature space
            dot_run1 = left_para.add_run('.' * 35)  # Signature line
            dot_run1.bold = False

            # Witness signature with proper spacing
            left_para.add_run('\n')  # 2 line breaks between sections
            left_para.add_run('Signature of witness:').bold = True
            left_para.add_run('\n\n')  # 3 line breaks for witness signature space
            dot_run2 = left_para.add_run('.' * 35)
            dot_run2.bold = False

            # Witness details with clean spacing
            left_para.add_run('\n')  # 2 line breaks before next field
            left_para.add_run('Name of witness:').bold = True
            left_para.add_run('\n\n')  # Single line break
            dot_run3 = left_para.add_run('.' * 35)
            dot_run3.bold = False

            left_para.add_run('\n')  # 2 line breaks before address
            left_para.add_run('Address of witness:').bold = True
            left_para.add_run('\n\n')  # Single line break
            dot_run4 = left_para.add_run('.' * 35)
            dot_run4.bold = False
            left_para.add_run('\n\n')  # Single line break
            dot_run5 = left_para.add_run('.' * 35)
            dot_run5.bold = False
            left_para.add_run('\n\n')  # Single line break
            dot_run6 = left_para.add_run('.' * 35)
            dot_run6.bold = False

            left_para.add_run('\n')  # 2 line breaks before occupation
            left_para.add_run('Occupation of witness:').bold = True
            left_para.add_run('\n\n')  # Single line break
            dot_run7 = left_para.add_run('.' * 35)
            dot_run7.bold = False

            # Right column - Second signatory (proper signature spacing)
            right_cell = sig_table.cell(0, 1)
            right_para = right_cell.paragraphs[0]
            right_para.paragraph_format.space_after = Pt(0)
            right_para.paragraph_format.line_spacing = 1.0

            # Main signature section with substantial space
            right_para.add_run('Signed and Delivered as a deed:').bold = True
            right_para.add_run('\n\n')  # 5 line breaks for signature space
            dot_run8 = right_para.add_run('.' * 35)  # Signature line
            dot_run8.bold = False

            # Witness signature with proper spacing
            right_para.add_run('\n')  # 2 line breaks between sections
            right_para.add_run('Signature of witness:').bold = True
            right_para.add_run('\n\n')  # 3 line breaks for witness signature space
            dot_run9 = right_para.add_run('.' * 35)
            dot_run9.bold = False

            # Witness details with clean spacing
            right_para.add_run('\n')  # 2 line breaks before next field
            right_para.add_run('Name of witness:').bold = True
            right_para.add_run('\n\n')  # Single line break
            dot_run10 = right_para.add_run('.' * 35)
            dot_run10.bold = False

            right_para.add_run('\n')  # 2 line breaks before address
            right_para.add_run('Address of witness:').bold = True
            right_para.add_run('\n\n')  # Single line break
            dot_run11 = right_para.add_run('.' * 35)
            dot_run11.bold = False
            right_para.add_run('\n\n')  # Single line break
            dot_run12 = right_para.add_run('.' * 35)
            dot_run12.bold = False
            right_para.add_run('\n\n')  # Single line break
            dot_run13 = right_para.add_run('.' * 35)
            dot_run13.bold = False

            right_para.add_run('\n')  # 2 line breaks before occupation
            right_para.add_run('Occupation of witness:').bold = True
            right_para.add_run('\n\n')  # Single line break
            dot_run14 = right_para.add_run('.' * 35)
            dot_run14.bold = False

            right_para.add_run('\n\n')  # 2 line breaks before note
            right_para.add_run('(*use a continuation sheet for additional signatories*)')

            # Add compact note at bottom
            doc.add_paragraph().paragraph_format.space_after = Pt(3)
            note = doc.add_paragraph()
            note_run = note.add_run(
                'Note - For execution and the attestation of the execution of a charge - see Rules 54 and 55.')
            note_run.bold = True
            note_run.font.size = Pt(9)
            note_run.font.name = 'Times New Roman'
            note.paragraph_format.space_after = Pt(0)

            # Save to BytesIO
            result = BytesIO()
            doc.save(result)
            result.seek(0)
            logger.info("Land Registry Form 51 created successfully")
            return result

        except Exception as e:
            logger.error(f"Error generating Land Registry Form 51: {str(e)}")
            import traceback
            logger.error(f"Full traceback: {traceback.format_exc()}")
            return None

    @classmethod
    def _add_form_row(cls, table, label, content, is_large=False, compact=False, fill_type='content'):
        """Add a row to the form table"""
        row = table.add_row()
        label_cell = row.cells[0]
        content_cell = row.cells[1]

        # Format label cell - more compact
        label_para = label_cell.paragraphs[0]
        label_para.paragraph_format.space_after = Pt(0)
        label_run = label_para.add_run(label)
        label_run.bold = True
        label_run.font.name = 'Times New Roman'
        label_run.font.size = Pt(10) if compact else Pt(11)
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Format content cell - more compact
        content_para = content_cell.paragraphs[0]
        content_para.paragraph_format.space_after = Pt(0)
        if compact:
            content_para.paragraph_format.line_spacing = 1.0

        if content:
            content_run = content_para.add_run(content)
            content_run.font.name = 'Times New Roman'
            content_run.font.size = Pt(10) if compact else Pt(11)
        elif fill_type == 'blank':
            # Leave completely blank for Date, Secured Party, Mortgagor
            pass
        else:
            # Add placeholder lines for other fields
            if is_large:
                content_run = content_para.add_run('_' * 70 + '\n' + '_' * 70 + '\n' + '_' * 70)
            else:
                content_run = content_para.add_run('_' * 60)
            content_run.font.name = 'Times New Roman'
            content_run.font.size = Pt(10) if compact else Pt(11)
            content_run.bold = False  # Make sure underscores are not bold

        content_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Set row height - more compact
        if is_large:
            row.height = Inches(0.8) if compact else Inches(1.2)
        else:
            row.height = Inches(0.4)  # Slightly more room for blank fields

    @classmethod
    def _serve_static_certificate_of_title(cls, requirement):
        """Serve the static Certificate of Title PDF file"""
        try:
            from django.conf import settings

            # Path to the static Certificate of Title PDF
            file_path = os.path.join(settings.BASE_DIR, 'static', 'documents', 'certificate_of_title.pdf')

            # Check if file exists
            if not os.path.exists(file_path):
                logger.error(f"Certificate of Title PDF not found at: {file_path}")
                return None

            # Read the file and return as BytesIO
            with open(file_path, 'rb') as file:
                file_content = file.read()

            result = BytesIO(file_content)
            result.seek(0)

            logger.info("Certificate of Title PDF served successfully")
            return result

        except Exception as e:
            logger.error(f"Error serving Certificate of Title PDF: {str(e)}")
            return None

    @classmethod
    def _get_template_context(cls, requirement):
        """Get context data for template rendering"""
        from django.conf import settings

        application = requirement.application
        deceased = application.deceased
        applicants = application.applicants.all()

        # Get today's date formatted
        today = timezone.now().strftime('%d %B %Y')

        # Get company info from settings and split address into lines
        company_address = getattr(settings, 'COMPANY_ADDRESS', '123 Main Street, City, Country')
        company_address_lines = [line.strip() for line in company_address.split(',')]

        company_info = {
            'name': getattr(settings, 'COMPANY_NAME', 'ALI Probate Ireland'),
            'address': company_address,
            'address_lines': company_address_lines
        }

        # Prepare solicitor info (user who created the application)
        solicitor_info = {
            'name': application.user.name if application.user else '',
            'address': None
        }

        if application.user and application.user.address:
            address = application.user.address
            solicitor_info['address'] = {
                'line1': address.line1,
                'line2': address.line2,
                'town_city': address.town_city,
                'county': address.county,
                'eircode': address.eircode
            }

        # Prepare application solicitor info (from application.solicitor field)
        application_solicitor_info = {
            'name': str(application.solicitor) if application.solicitor else '____________'
        }

        # Prepare deceased info
        deceased_info = {
            'full_name': f"{deceased.first_name} {deceased.last_name}" if deceased else '',
            'name_with_deceased': f"{deceased.first_name} {deceased.last_name} (Deceased)" if deceased else '____________ (Deceased)'
        }

        # Prepare beneficiaries/applicants info
        beneficiaries = []
        for applicant in applicants:
            beneficiaries.append({
                'first_name': applicant.first_name,
                'last_name': applicant.last_name,
                'full_name': f"{applicant.first_name} {applicant.last_name}"
            })

        # Get beneficiaries names string
        beneficiaries_names = " & ".join([b['full_name'] for b in beneficiaries]) if beneficiaries else '____________'

        # Get estate value
        estate_value = application.value_of_the_estate_after_expenses()

        return {
            'today': today,
            'application': application,
            'deceased': deceased_info,
            'solicitor': solicitor_info,
            'application_solicitor': application_solicitor_info,
            'company': company_info,
            'beneficiaries': beneficiaries,
            'beneficiaries_names': beneficiaries_names,
            'loan_application_number': application.id,
            'currency_sign': application.user.get_currency() if application.user else '€',
            'estate_value': estate_value,
        }

    @classmethod
    def _get_land_registry_context(cls, requirement):
        """Get context data specifically for Land Registry Form 51"""
        application = requirement.application

        # Basic context - can be expanded later with actual data
        context = {
            'date': '',
            'secured_party': '',
            'mortgagor': '',
            'folio_number': '',
            'county': '',
            'property_address': '',
        }

        # Future enhancement: populate from application data
        # Example:
        # if hasattr(application, 'property') and application.property:
        #     context['property_address'] = application.property.address
        #     context['county'] = application.property.county
        # if hasattr(application, 'borrower') and application.borrower:
        #     context['mortgagor'] = application.borrower.full_name

        return context

    @classmethod
    def get_filename(cls, requirement):
        """Get the filename for the generated document"""
        application = requirement.application

        if "Beneficiaries Irrevocable Instruction to Law Firm" in requirement.document_type.name:
            return f"Beneficiaries_Irrevocable_Instruction_to_Law_Firm_{application.id}.pdf"
        elif "Solicitor Letter of Undertaking (Not to distribute estate)" in requirement.document_type.name:
            return f"Solicitor_Letter_of_Undertaking_{application.id}.pdf"
        elif "Renunciation of Probate" in requirement.document_type.name:
            return f"Renunciation_of_Probate_{application.id}.docx"
        elif "Certificate of Title" in requirement.document_type.name:
            return f"Certificate_of_Title.pdf"
        elif "Land Registry Form 51" in requirement.document_type.name:
            return f"Land_Registry_Form_51_{application.id}.docx"
        else:
            return f"{requirement.document_type.name}_{application.id}.pdf"

    @classmethod
    def get_content_type(cls, requirement):
        """Get the content type for the generated document"""
        if ("Renunciation of Probate" in requirement.document_type.name or
                "Land Registry Form 51" in requirement.document_type.name):
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            return "application/pdf"

    # Legacy method for backward compatibility
    @classmethod
    def generate_pdf_response(cls, requirement):
        """Legacy method - redirects to generate_document_response"""
        return cls.generate_document_response(requirement)
